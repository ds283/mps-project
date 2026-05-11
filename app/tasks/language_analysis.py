#
# Created by David Seery on 03/04/2026.
# Copyright (c) 2026 University of Sussex. All rights reserved.
#
# This file is part of the MPS-Project platform developed in
# the School of Mathematics & Physical Sciences, University of Sussex.
#
# Contributors: David Seery <D.Seery@sussex.ac.uk>
#

import hashlib
import re
import time
import unicodedata

import numpy as np
from celery import chord, states
from celery import group as cgroup
from flask import current_app
from sqlalchemy.exc import SQLAlchemyError

from ..database import db
from ..models import (
    ProjectClass,
    ProjectClassConfig,
    SubmissionPeriodRecord,
    SubmissionRecord,
    TaskRecord,
    Tenant,
)
from ..shared.ai_calibration import mahalanobis_distance
from ..shared.asset_tools import AssetCloudAdapter
from ..shared.llm_services import _TOKENS_PER_WORD, _call_llm, _truncate_text

# Tokens-per-word estimate for student submission content.  Technical/academic text with
# equations, DOIs, code snippets, and jargon.  Empirical calibration (comparing Ollama-reported
# prompt tokens against est_tok across many submissions) shows ~1.44 t/w actual; 1.5 provides a
# ~4% safety margin.  Used in word-budget division for chunk sizing and in est_tok computation
# via the user_tokens_per_word argument to _call_llm.
_TOKENS_PER_WORD_CONTENT = 1.5
from ..shared.llm_thresholds import (
    classify_burstiness,
    classify_mattr,
    classify_mtld,
    classify_sentence_cv,
)
from ..shared.scraped_text_store import get_scraped_text, store_scraped_text
from ..shared.text_utils import (
    _APPENDIX_HEADING,
    _split_document,
    _strip_math_lines,
    _strip_toc_lines,
)
from ..shared.workflow_logging import log_db_commit
from ..task_queue import progress_update
from .pipeline_tracking import get_pipeline_redis, record_step_end, record_step_start

# ---------------------------------------------------------------------------
# Prompt versioning.
# Bump PROMPT_VERSION whenever prompt text is intentionally changed, so stored
# results can be compared against the prompt that generated them.
#
# MAINTENANCE NOTE — standalone script mirror:
# This file is mirrored by lexical-pipeline-validation/language_analysis_core.py
# for use in offline analysis. Any changes to metric computation (MATTR, MTLD,
# sentence_cv, burstiness, reference counting, LLM prompt text) must be reflected
# there. Related standalone scripts that may also need updating:
#   - lexical-pipeline-validation/arxiv_control_analysis.py  (cache + metric collection)
#   - lexical-pipeline-validation/analysis_cache.py          (SQLite cache schema)
#   - lexical-pipeline-validation/lexical_diversity_pipeline.py (Mahalanobis / calibration)
#
# PROMPT_VERSION (here) and PIPELINE_VERSION in language_analysis_core.py should
# be kept in sync: bump both when the LLM prompt text changes intentionally.
# ---------------------------------------------------------------------------

PROMPT_VERSION = 1

# Bump STATS_ALGORITHM_VERSION when the compute_statistics logic changes in a way
# that would produce materially different output from the same input text.
STATS_ALGORITHM_VERSION = 1


def _prompt_hash(prompt: str) -> str:
    return hashlib.sha256(prompt.encode()).hexdigest()[:12]


# ---------------------------------------------------------------------------
# Pre-loaded rubric snapshot, safe to use after db.session.close().
# ---------------------------------------------------------------------------


class _RubricSnapshot:
    """Holds eagerly-loaded rubric data extracted from the ORM before session close."""

    __slots__ = ("id", "label", "_bands")

    def __init__(self, rubric):
        self.id = rubric.id
        self.label = rubric.label
        self._bands = [
            {
                "id": b.id,
                "label": b.label,
                "criteria": [
                    {"id": c.id, "text": c.text, "tag": c.tag} for c in b.criteria
                ],
            }
            for b in rubric.bands
        ]

    def to_prompt_bands(self):
        return [
            {"band": b["label"], "criteria": [c["text"] for c in b["criteria"]]}
            for b in self._bands
        ]

    def negative_criteria(self):
        return frozenset(
            c["text"]
            for b in self._bands
            for c in b["criteria"]
            if c["tag"] == "negative"
        )

    def positive_floor_criteria(self):
        return frozenset(
            c["text"]
            for b in self._bands
            for c in b["criteria"]
            if c["tag"] == "positive_floor"
        )


# ---------------------------------------------------------------------------
# Word groups for burstiness analysis.
# Each group defines a set of canonical spaCy lemmas. The corpus is
# lemmatised and token positions are matched against these lemma sets.
# ---------------------------------------------------------------------------

BURSTINESS_WORD_GROUPS = [
    {"group": "suggest", "lemmas": {"suggest"}},
    {"group": "indicate", "lemmas": {"indicate"}},
    {"group": "demonstrate", "lemmas": {"demonstrate"}},
    {"group": "show", "lemmas": {"show"}},
    {"group": "appear", "lemmas": {"appear"}},
    {"group": "estimate", "lemmas": {"estimate"}},
    {"group": "assume", "lemmas": {"assume"}},
    {"group": "imply", "lemmas": {"imply"}},
    {"group": "significant", "lemmas": {"significant", "insignificant"}},
    {"group": "important", "lemmas": {"important", "relevant"}},
    {"group": "consistent", "lemmas": {"consistent", "inconsistent"}},
    {"group": "unexpected", "lemmas": {"unexpected", "surprising"}},
    {"group": "clear", "lemmas": {"clear", "unclear"}},
    {"group": "compare", "lemmas": {"compare"}},
    {"group": "contrast", "lemmas": {"contrast", "differ"}},
    {"group": "agree", "lemmas": {"agree", "disagree", "confirm"}},
    {"group": "support", "lemmas": {"support", "contradict"}},
]

# Minimum number of occurrences for a word group to be included in the
# burstiness calculation. Groups below this threshold are excluded
# because the inter-arrival distribution is too sparse.
BURSTINESS_MIN_OCCURRENCES = 8


# ---------------------------------------------------------------------------
# Hedging and filler patterns for AI-tendency detection.
# These are matched case-insensitively.
# ---------------------------------------------------------------------------

HEDGING_PATTERNS = [
    r"it is important to note that",
    r"it is worth noting that",
    r"it is crucial to note that",
    r"it should be noted that",
    r"needless to say",
    r"it goes without saying that",
    r"as mentioned above",
    r"as noted above",
    r"as discussed previously",
    r"it is interesting to note that",
    r"\bsignificantly\b",
    r"\bimportantly\b",
    r"\bfundamentally\b",
    r"\bessentially\b",
    r"\bbasically\b",
    r"it is clear that",
    r"\bclearly\b",
    r"\bobviously\b",
    r"\bevidently\b",
    r"of course",
    r"\bnaturally\b",
]

FILLER_PATTERNS = [
    r"\bfurthermore\b",
    r"\bmoreover\b",
    r"in conclusion",
]

# Unicode em-dash and ASCII triple-hyphen (sometimes produced by PDF extraction
# from LaTeX --- ligature)
EM_DASH_PATTERN = r"(\u2014|---)"


# ---------------------------------------------------------------------------
# Lazy spaCy model loader.
# The model is loaded once per worker process on first use.
# ---------------------------------------------------------------------------

_nlp = None


def _get_nlp():
    global _nlp
    if _nlp is None:
        import spacy

        _nlp = spacy.load("en_core_web_sm", disable=["ner", "parser"])
        # The parser normally provides sentence boundaries; since it is disabled,
        # add the rule-based sentencizer so that doc.sents is available.
        if not _nlp.has_pipe("senter") and not _nlp.has_pipe("sentencizer"):
            _nlp.add_pipe("sentencizer", first=True)
    return _nlp


# ---------------------------------------------------------------------------
# Text extraction helpers.
# ---------------------------------------------------------------------------


def _extract_pdf_text(path: str) -> tuple[str, int]:
    """
    Extract plain text from a PDF file at *path* using PyMuPDF.

    Returns (raw_text, page_count).  Header and footer regions (top/bottom 8%
    of each page) are skipped to reduce noise from running headers and page
    numbers.
    """
    import fitz

    doc = fitz.open(path)
    pages = []
    for page in doc:
        page_height = page.rect.height
        blocks = page.get_text("blocks")  # (x0, y0, x1, y1, text, block_no, block_type)
        body_blocks = [
            b[4]
            for b in blocks
            if b[1] > page_height * 0.08  # skip top 8 % (header / running title)
            and b[3] < page_height * 0.92  # skip bottom 8 % (footer / page number)
            and b[6] == 0  # block_type 0 = text (not image)
        ]
        pages.append("\n".join(body_blocks))
    page_count = len(doc)
    doc.close()
    return "\n\n".join(pages), page_count


def _extract_docx_text(path: str) -> tuple[str, int]:
    """
    Extract plain text from a Word document at *path* using python-docx.

    This is a basic implementation: it joins all paragraph texts.
    Table cell content is included.  Returns (raw_text, page_count=0) since
    Word documents do not expose a reliable page count without rendering.
    """
    try:
        from docx import Document

        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        return "\n\n".join(paragraphs), 0
    except Exception as exc:
        current_app.logger.warning(f"language_analysis: DOCX extraction failed: {exc}")
        return "", 0


# ---------------------------------------------------------------------------
# Statistical analysis helpers.
# ---------------------------------------------------------------------------

# Caption line pattern (figure or table captions, used to exclude word count).
# Handles simple (Figure 1), chapter-relative (Figure 3.1), and appendix (Figure A.1)
# numbering styles.
_CAPTION_LINE = re.compile(
    r"^\s*(?:figure|fig\.|table)\s+(?:[A-Z]\.)?\d+(?:\.\d+)?",
    re.IGNORECASE | re.MULTILINE,
)

# Figure/table reference label patterns.
#
# group(1) — canonical label key: "1", "3.1", "A.1"
# group(2) — optional subfigure suffix letter ("a", "b") — NOT included in the key.
#
# (?:[A-Z]\.)? with IGNORECASE matches appendix-letter prefixes ("A.", "B.") but not
# digit characters, so "3." in "3.1" is never mistaken for a prefix.
_FIG_REF = re.compile(
    r"\b(?:figure|fig\.)\s+((?:[A-Z]\.)?\d+(?:\.\d+)?)([a-z])?",
    re.IGNORECASE,
)
_TAB_REF = re.compile(
    r"\btable\s+((?:[A-Z]\.)?\d+(?:\.\d+)?)([a-z])?",
    re.IGNORECASE,
)

# Numbered bibliography entry: [N] or N. followed by an uppercase letter.
# \d{1,3} rejects large numbers (page ranges, ISSN fragments) that can appear at
# line-starts after PDF text wrapping.  Requiring [A-Z] after the whitespace rules
# out "3313.\n  issn:…" continuations; real entries always start with an author
# surname or institution name (uppercase).
_NUMBERED_ENTRY = re.compile(r"^\s*(\[\d{1,3}\]|\d{1,3}\.)\s+[A-Z]", re.MULTILINE)

# Numbered citation in main text: [N] or [N, M, ...]
_NUMBERED_CITATION = re.compile(r"\[(\d+(?:\s*,\s*\d+)*)\]")

# Author-year reference entry signals.
# A year in parentheses is the minimal marker of an author-year bibliography entry.
_REF_YEAR = re.compile(r"\(\d{4}[a-z]?\)")
# arXiv identifier — near-certain evidence that a line belongs to the reference list.
_ARXIV_ID = re.compile(r"\barXiv:\d{4}\.\d{4,5}\b", re.IGNORECASE)
# DOI — also strong evidence of a reference list line.
_DOI = re.compile(r"\bdoi:\s*10\.\d{4,}", re.IGNORECASE)


# Code-listing detection used to exclude source-code sentences from sentence CV.
_CODE_CHARS = frozenset("=()[]{}#")
_CODE_CHAR_RATIO_THRESHOLD = 0.04  # >4 % of chars are code punctuation
_UNDERSCORE_TOKEN_THRESHOLD = 0.15  # >15 % of whitespace-split tokens contain '_'


_MIN_CODE_BLOCK_LINES = 3  # runs shorter than this are kept (false-positive protection)


def _strip_code_blocks(text: str, min_run: int = _MIN_CODE_BLOCK_LINES) -> str:
    """Remove contiguous blocks of source-code lines from *text*.

    Uses _looks_like_code() per line, but only strips lines that belong to a
    run of >= min_run consecutive code-like lines.  Blank lines inside a run
    are treated as neutral and do not break it — a blank line between two code
    lines stays in the run, matching how code is commonly extracted from PDFs.

    Mirrors strip_code_blocks() in language_analysis_core.py.
    """
    lines = text.splitlines()

    # Tag each line: True = code-like, False = prose, None = blank
    tags: list[bool | None] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            tags.append(None)
        else:
            tags.append(_looks_like_code(stripped))

    # Resolve blank lines: a blank line is treated as code only when both its
    # preceding and following non-blank lines are code-like.
    # Forward pass: tentatively propagate True through blanks.
    resolved: list[bool] = [False] * len(tags)
    last = False
    for i, t in enumerate(tags):
        if t is True:
            last = True
        elif t is False:
            last = False
        else:  # blank
            pass  # keep last unchanged
        resolved[i] = last

    # Backward pass: un-mark blank lines whose next non-blank line is prose.
    last = False
    for i in range(len(tags) - 1, -1, -1):
        t = tags[i]
        if t is True:
            last = True
        elif t is False:
            last = False
        # blank inherits from forward pass unless backward pass disagrees
        if tags[i] is None and not last:
            resolved[i] = False

    # Remove runs of True that are long enough.
    keep = [True] * len(lines)
    i = 0
    while i < len(resolved):
        if resolved[i]:
            j = i
            while j < len(resolved) and resolved[j]:
                j += 1
            if (j - i) >= min_run:
                for k in range(i, j):
                    keep[k] = False
            i = j
        else:
            i += 1

    return "\n".join(line for line, k in zip(lines, keep) if k)


def _word_count(main_text: str) -> int:
    """
    Count words in *main_text*, excluding caption lines.
    This is an estimate; punctuation is included with adjacent words.
    """
    clean = _CAPTION_LINE.sub("", main_text)
    return len(clean.split())


def _count_bibliography(biblio_text: str) -> tuple[int, list[str]]:
    """
    Count bibliography entries and return (count, list_of_entry_keys).

    Handles:
    - Numbered entries starting with [N] or N.
    - Falls back to a heuristic line count for author-year styles.

    For the heuristic fallback the text is first truncated at any appendix
    heading (appendices frequently follow the reference list and would
    otherwise inflate the count).  Lines are then filtered to those that
    carry strong signals of an author-year entry: a year in parentheses,
    an arXiv identifier, or a DOI.  If fewer than three such candidate
    lines are found the simple total-line count is used as a fallback to
    avoid under-counting documents whose format is unusual.
    """
    # Trim at the first appendix heading so appendix content is not counted.
    app_match = _APPENDIX_HEADING.search(biblio_text)
    if app_match:
        biblio_text = biblio_text[: app_match.start()]

    numbered = _NUMBERED_ENTRY.findall(biblio_text)
    if numbered:
        # Strip to just the bracket/dot keys for cross-referencing
        keys = [k.strip().strip(".").strip("[]") for k in numbered]
        # Validate: a genuine numbered bibliography must start near 1.
        # If the minimum numeric key is large (e.g. 82 from a wrapped page-range
        # "p.\n82. TITLE") treat the whole match as a false positive and fall through
        # to the author-year heuristic.
        try:
            if min(int(k) for k in keys) <= 5:
                return len(keys), keys
        except ValueError:
            return len(keys), keys  # non-integer keys — trust the match

    # Heuristic: count non-blank lines after the heading as a rough entry count.
    lines = [ln.strip() for ln in biblio_text.splitlines() if ln.strip()]
    # Skip the heading line itself.
    body_lines = [
        ln
        for ln in lines[1:]
        if ln and not ln.lower().startswith(("ref", "biblio", "works"))
    ]

    # Stage 1: prefer lines that look like author-year reference entries.
    # Signals: year in parentheses, arXiv identifier, or a DOI.
    entry_candidates = [
        ln
        for ln in body_lines
        if _REF_YEAR.search(ln) or _ARXIV_ID.search(ln) or _DOI.search(ln)
    ]

    # Stage 2: if too few candidates found, fall back to the total line count.
    if len(entry_candidates) >= 3:
        return len(entry_candidates), []
    return len(body_lines), []


def _check_uncited(main_text: str, keys: list[str]) -> list[str]:
    """
    For a numbered bibliography with *keys* like ['1', '2', ...], return
    those keys that do not appear as citations in *main_text*.
    Only meaningful for numbered (LaTeX-style) references.
    """
    if not keys:
        return []
    cited_numbers: set[str] = set()
    for match in _NUMBERED_CITATION.finditer(main_text):
        for num in match.group(1).split(","):
            cited_numbers.add(num.strip())
    return [k for k in keys if k not in cited_numbers]


def _check_figure_table_refs(text: str) -> tuple[list[str], list[str]]:
    """
    Return (uncaptioned_figures, uncaptioned_tables): labels that appear in
    captions but are not mentioned anywhere in the main body text.

    Labels are normalised to their canonical form before deduplication:
    - Chapter-relative numbers: "Figure 3.1", "Figure 4.5" → keys "3.1", "4.5"
    - Appendix figures:         "Figure A.1", "Figure B.3" → keys "A.1", "B.3"
    - Subfigure panels:         "Figure 3.1a", "Figure 3.1b" → key "3.1"
      (panel letters are stripped; panels are not counted as separate figures)
    """
    fig_labels: dict[str, str] = {}  # key -> display string, e.g. "3.1" -> "Figure 3.1"
    tab_labels: dict[str, str] = {}

    # _FIG_REF group(1) = canonical key; group(2) = optional subfigure suffix (ignored).
    for m in _FIG_REF.finditer(text):
        n = m.group(1)
        if n not in fig_labels:
            fig_labels[n] = f"Figure {n}"

    for m in _TAB_REF.finditer(text):
        n = m.group(1)
        if n not in tab_labels:
            tab_labels[n] = f"Table {n}"

    # A label is considered cited if it appears more than once in the text.
    # The citation pattern accepts an optional trailing subfigure suffix so that
    # "Figure 3.1a" counts as a citation of key "3.1".
    def _cited(pattern: str) -> bool:
        return len(re.findall(pattern, text, re.IGNORECASE)) > 1

    uncaptioned_figs = [
        label
        for n, label in fig_labels.items()
        if not _cited(rf"\b(?:figure|fig\.)\s+{re.escape(n)}[a-z]?\b")
    ]
    uncaptioned_tabs = [
        label
        for n, label in tab_labels.items()
        if not _cited(rf"\btable\s+{re.escape(n)}[a-z]?\b")
    ]

    return uncaptioned_figs, uncaptioned_tabs


def _compute_mattr_mtld(main_text: str) -> tuple[float | None, float | None]:
    """
    Compute MATTR (window=100) and MTLD for *main_text*.
    Returns (mattr, mtld), either of which may be None on failure.
    Requires at least 100 words; fewer returns (None, None).

    Code blocks are stripped before analysis to prevent source-code identifiers
    from artificially inflating vocabulary diversity.
    """
    try:
        from lexicalrichness import LexicalRichness

        lex = LexicalRichness(_strip_code_blocks(main_text))
        if lex.words < 100:
            print(
                f"_compute_matr_mtld: too few words to compute MATTR and MTLD statistics ({lex.words} words detected)"
            )
            return None, None
        mattr = float(lex.mattr(window_size=100))
        mtld = float(lex.mtld(threshold=0.72))
        return mattr, mtld
    except Exception as exc:
        current_app.logger.warning(
            f"language_analysis: MATTR/MTLD computation failed: {exc}"
        )
        return None, None


def _compute_burstiness(text: str) -> tuple[dict, float | None]:
    """
    Compute per-group and aggregate Goh-Barabási burstiness for *text*.

    Returns (group_results_dict, aggregate_burstiness).
    *group_results_dict* maps group name -> B value (or None if excluded).
    *aggregate_burstiness* is the mean over eligible groups, or None.
    """
    nlp = _get_nlp()

    # spaCy processes the text; we only need token lemmas and positions.
    # Restrict to alphabetic tokens to avoid punctuation noise.
    doc = nlp(text)
    token_lemmas = [
        (i, token.lemma_.lower()) for i, token in enumerate(doc) if token.is_alpha
    ]

    group_results: dict[str, float | None] = {}
    eligible_values: list[float] = []

    for group_def in BURSTINESS_WORD_GROUPS:
        group_name = group_def["group"]
        target_lemmas: set[str] = group_def["lemmas"]

        positions = [i for i, lemma in token_lemmas if lemma in target_lemmas]

        if len(positions) < BURSTINESS_MIN_OCCURRENCES:
            group_results[group_name] = None
            continue

        arrivals = np.diff(positions).astype(float)
        mu = float(np.mean(arrivals))
        sigma = float(np.std(arrivals, ddof=1)) if len(arrivals) > 1 else 0.0

        denom = sigma + mu
        B = float((sigma - mu) / denom) if denom != 0 else 0.0

        group_results[group_name] = B
        eligible_values.append(B)

    aggregate = float(np.mean(eligible_values)) if eligible_values else None
    return group_results, aggregate


def _looks_like_code(text: str) -> bool:
    """Return True when *text* looks more like source code than prose.

    Two independent signals, either of which is sufficient:
      1. High density of code-typical punctuation (=, (, ), [, ], {, }, #).
      2. High fraction of whitespace-split tokens that contain an underscore
         (Python/C-style identifiers).
    """
    if not text:
        return False
    code_ratio = sum(1 for ch in text if ch in _CODE_CHARS) / len(text)
    if code_ratio > _CODE_CHAR_RATIO_THRESHOLD:
        return True
    tokens = text.split()
    if tokens:
        underscore_ratio = sum(1 for t in tokens if "_" in t) / len(tokens)
        if underscore_ratio > _UNDERSCORE_TOKEN_THRESHOLD:
            return True
    return False


def _compute_sentence_cv(text: str) -> float | None:
    """
    Compute the coefficient of variation (CV = σ/μ) of sentence lengths for *text*.

    Sentence length is measured as the number of non-punctuation, non-space tokens
    per sentence, using spaCy's sentence segmentation.  Returns None if fewer than
    5 sentences are found (insufficient data for a stable estimate).

    Sentences that look like source code (high code-punctuation density or high
    underscore-identifier fraction) are excluded before computing the CV.
    """
    nlp = _get_nlp()
    doc = nlp(text)

    lengths = [
        sum(1 for tok in sent if not tok.is_punct and not tok.is_space)
        for sent in doc.sents
        if not _looks_like_code(sent.text)
    ]
    # Discard empty or single-token sentences (e.g. section headings misread as sentences)
    lengths = [ln for ln in lengths if ln > 1]

    if len(lengths) < 5:
        return None

    mu = float(np.mean(lengths))
    if mu == 0.0:
        return None
    sigma = float(np.std(lengths, ddof=1))
    return sigma / mu


def _count_patterns(text: str) -> dict:
    """
    Count occurrences of AI-tendency indicator patterns.
    Returns a dict with per-pattern counts and totals.
    """
    hedging_total = 0
    hedging_detail: dict[str, int] = {}
    for pattern in HEDGING_PATTERNS:
        count = len(re.findall(pattern, text, re.IGNORECASE))
        hedging_detail[pattern] = count
        hedging_total += count

    filler_total = 0
    filler_detail: dict[str, int] = {}
    for pattern in FILLER_PATTERNS:
        count = len(re.findall(pattern, text, re.IGNORECASE))
        filler_detail[pattern] = count
        filler_total += count

    em_dash_count = len(re.findall(EM_DASH_PATTERN, text))

    return {
        "hedging_total": hedging_total,
        "hedging_detail": hedging_detail,
        "filler_total": filler_total,
        "filler_detail": filler_detail,
        "em_dash_count": em_dash_count,
    }


def _ai_concern_flag(
    mattr: float | None,
    mtld: float | None,
    sentence_cv: float | None,
    calibrations: list | None,
    mean_nll: float | None = None,
    nll_cv: float | None = None,
    llm_model_name: str | None = None,
    llm_context_window: int | None = None,
) -> dict:
    """
    Classify the overall AI concern level using Mahalanobis distance tests
    across all applicable TenantAICalibration objects.

    For each calibration:
      "lexical" (3D) — uses (MATTR, MTLD, sentence_cv) if all are available.
      "full"    (5D) — uses (MATTR, MTLD, sentence_cv, mean_nll, nll_cv) if all
                       are available and the calibration's LLM model/context window
                       matches (llm_model_name, llm_context_window).
                       Legacy 4D "full" calibrations use only the first 4 features.

    Bonferroni correction: per-test alpha = 0.05/K (medium) and 0.01/K (high),
    where K is the number of calibrations actually evaluated.  The flag fires
    if any individual test exceeds its corrected threshold.

    Graceful degradation
    --------------------
    Returns concern="uncalibrated" (sigma=None, p_value=None) when:
      - calibrations is None or empty, or
      - no calibration has all required features available.

    Returns
    -------
    dict with keys:
        "concern"  : "low" | "medium" | "high" | "uncalibrated"
        "sigma"    : float | None  — sigma from the most significant test
        "p_value"  : float | None  — p_value from the most significant test
    """
    _UNCALIBRATED = {"concern": "uncalibrated", "sigma": None, "p_value": None}

    if not calibrations:
        return _UNCALIBRATED

    # Build the list of (calibration_obj, feature_vector) pairs that can be
    # evaluated given the metrics available for this submission.
    applicable = []
    for cal in calibrations:
        if cal.feature_set == "lexical":
            if mattr is None or mtld is None or sentence_cv is None:
                continue
            applicable.append((cal, [mattr, mtld, sentence_cv]))
        elif cal.feature_set == "full":
            if not cal.is_llm_matched(llm_model_name, llm_context_window):
                continue
            # Slice to cal.n_features so legacy 4D calibrations only require mean_nll.
            full_features = [mattr, mtld, sentence_cv, mean_nll, nll_cv]
            features = full_features[: cal.n_features]
            if any(v is None for v in features):
                continue
            applicable.append((cal, features))

    K = len(applicable)
    if K == 0:
        return _UNCALIBRATED

    alpha_medium = 0.05 / K
    alpha_high = 0.01 / K

    best_sigma: float | None = None
    best_p_value: float = 1.0
    overall_concern = "low"
    cal_results = []

    for cal, features in applicable:
        try:
            sigma, p_value = mahalanobis_distance(features, cal)
        except Exception:
            continue

        if best_sigma is None or p_value < best_p_value:
            best_sigma = sigma
            best_p_value = p_value

        if p_value <= alpha_high:
            this_concern = "high"
            overall_concern = "high"
        elif p_value <= alpha_medium:
            this_concern = "medium"
            if overall_concern != "high":
                overall_concern = "medium"
        else:
            this_concern = "low"

        cal_results.append(
            {
                "feature_set": cal.feature_set,
                "n_features": cal.n_features,
                "llm_model_name": cal.llm_model_name,
                "llm_context_window": cal.llm_context_window,
                "sigma": sigma,
                "p_value": p_value,
                "concern": this_concern,
            }
        )

    if best_sigma is None:
        return _UNCALIBRATED

    return {
        "concern": overall_concern,
        "sigma": best_sigma,
        "p_value": best_p_value,
        "calibration_results": cal_results,
        "bonferroni_k": K,
        "bonferroni_alpha_medium": alpha_medium,
        "bonferroni_alpha_high": alpha_high,
    }


# ---------------------------------------------------------------------------
# LLM helpers.
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Context-window-aware chunking constants.
# ---------------------------------------------------------------------------

# Minimum context window the synthesis pass requires.  Ollama must be
# configured with num_ctx ≥ this value (and ≥ OLLAMA_CONTEXT_SIZE).
_SYNTHESIS_MIN_CTX = 8192

# Overhead tokens for the dedicated metadata extraction call:
#   small system prompt (~250) + small schema response (~250).
_METADATA_OVERHEAD_TOKENS = 500

# Overhead tokens for a single feedback chunk call:
#   small system prompt (~200) + feedback array response (~200).
_FEEDBACK_OVERHEAD_TOKENS = 400

# Maximum evidence entries retained per criterion code after map-phase
# aggregation.  Minority-polarity entries are always kept (one each).
_MAX_EVIDENCE_PER_CRITERION = 3

# Excerpt character limit applied when the synthesis evidence text would
# otherwise overflow the synthesis context window.
_MAX_EXCERPT_CHARS = 150


def _build_system_prompt(truncated: bool, rubric) -> str:
    """Construct the LLM system prompt including the rubric and explicit criteria checklist."""
    prompt_bands = rubric.to_prompt_bands()
    negative_criteria = rubric.negative_criteria()
    positive_floor_criteria = rubric.positive_floor_criteria()

    # Rubric narrative section (used for context and assessment rules)
    band_sections = []
    for band in prompt_bands:
        criteria_list = "\n".join(f"  - {c}" for c in band["criteria"])
        band_sections.append(f"### {band['band']}\n{criteria_list}")
    rubric_text = "\n\n".join(band_sections)

    # Explicit numbered checklist — enumerate every criterion so the LLM cannot
    # accidentally omit or paraphrase any of them.  The criterion text here must
    # be reproduced verbatim in the "criterion" field of each response entry.
    checklist_sections = []
    for band_idx, band in enumerate(prompt_bands, start=1):
        lines = [f"Band {band_idx} — {band['band']}:"]
        for crit_idx, criterion in enumerate(band["criteria"], start=1):
            if criterion in negative_criteria:
                tag = "  [NEGATIVE — 'Not evident' for submissions above this band]"
            elif criterion in positive_floor_criteria:
                tag = "  [POSITIVE FLOOR — 'Strong evidence' for submissions above this band]"
            else:
                tag = ""
            lines.append(f"  {band_idx}.{crit_idx}  {criterion}{tag}")
        checklist_sections.append("\n".join(lines))
    criteria_checklist = "\n\n".join(checklist_sections)

    truncation_note = (
        "\n\nNOTE: The document text provided may have been truncated. "
        "If you see the marker '[... middle section omitted due to length ...]', "
        "the middle portion of the document has been removed. Please note this "
        "limitation in the caveats field of your response."
        if truncated
        else ""
    )

    schema_description = f"""
Your response must be a JSON object with the following fields:

- "report_summary": 1-2 paragraphs describing what the report is about, what the student
  did, and the overall approach taken. This is a content summary, not a marking judgement.

- "classification": the recommended grade band (one of the exact band names from the rubric).

- "overall_reasoning": one paragraph explaining why you recommend that band.

- "bands": an array with exactly one entry per grade band, in order from lowest to highest.
  Each entry must have:
    - "band": the band name, copied exactly from the rubric
    - "band_assessment": one of:
        "Criteria clearly demonstrated"   — report substantially satisfies all criteria for this band
        "Criteria mostly demonstrated"    — most criteria met, minor gaps
        "Some criteria demonstrated"      — meaningful evidence for some criteria but threshold not reached
        "No evidence"                     — report does not meet this band's criteria
    - "criteria": an object with exactly one key per criterion for that band.
      Each key is the short numeric code from the checklist below (e.g. "1.1", "2.3").
      Each value is an object with:
        - "assessment": one of "Strong evidence", "Partial evidence", or "Not evident"
        - "commentary": one or two sentences with specific textual evidence from the report
        - "confidence": one of "high", "medium", or "low"
      Every code in the checklist MUST appear as a key. Do not add extra keys.

- "caveats": any limitations, uncertainties, or criteria that could not be assessed from
  text alone (e.g. typesetting quality, visual presentation).

## Criteria checklist — use the numeric code (e.g. "1.1") as the key for each criterion

{criteria_checklist}

Flag criteria that relate to visual presentation or typesetting with low confidence and a note
that these cannot be fully assessed from extracted text alone.
"""

    cross_band_guidance = """
## Cross-band assessment guidance

Rubric criteria fall into two types:

1. **Negative/limitation criteria** describe deficiencies (e.g. "Scientific work of \
limited quality", "Limited evidence for technical and practical skills"). For a \
submission that exceeds that band, the correct assessment is "Not evident" — the \
deficiency is simply absent.

2. **Positive minimum criteria** describe a floor requirement (e.g. "At least some \
attempt to explain and interpret the results", "Some discussion of future directions"). \
A submission at a higher band clearly meets and exceeds these floors. The correct \
assessment is "Strong evidence" — not "Not evident".

For bands BELOW your recommended classification:
- Negative/limitation criteria → "Not evident"
- Positive minimum criteria → "Strong evidence"

For your recommended band: assess each criterion normally on the evidence in the text.

For bands ABOVE your recommended classification: use "Partial evidence" or "Not evident" \
with commentary explaining specifically what evidence is missing or insufficient.

Your assessments must be self-consistent. If you recommend "2.1 class", marking \
"At least some attempt to explain and interpret the results" as "Not evident" is \
contradictory — a 2.1 submission clearly demonstrates this.
"""

    return f"""You are an expert academic marker assisting with assessment of undergraduate and postgraduate project reports. Your task is to evaluate the submitted report text against the rubric below and recommend a grade band.

## Rubric

{rubric_text}

## Assessment Rules

Each higher grade band subsumes all criteria of lower bands. A submission cannot be awarded a higher band unless it substantially satisfies the criteria of all lower bands. Assess against the highest band whose full set of criteria, including those of all lower bands, are mostly satisfied.

Where you are uncertain, or where the text does not provide enough evidence, say so explicitly. Quote specific passages from the text to support your assessment — this makes your reasoning auditable by a human marker. Flag criteria where you consider yourself less able to assess reliably from text alone.
{cross_band_guidance}
{schema_description}{truncation_note}"""


def _build_user_prompt(document_text: str) -> str:
    return f"Please assess the following report text:\n\n---\n\n{document_text}\n\n---"


_REQUIRED_JSON_KEYS = {
    "report_summary",
    "classification",
    "overall_reasoning",
    "bands",
    "caveats",
}


def _make_band_schema(band_idx: int, criteria_list: list) -> dict:
    """
    Return the JSON Schema for a single band entry within the top-level 'bands' object.

    *band_idx* is the 1-based index of the band (used to generate criterion codes).
    *criteria_list* is the exhaustive list of criterion strings for this band.

    'criteria' is modelled as a fixed-key JSON Schema object (not an array) where each
    property key is a short numeric code ("1.1", "1.2", …) rather than the full criterion
    text.  Using short codes avoids a SIGSEGV in the llama.cpp GBNF grammar generator,
    which crashes when property name literals in the generated grammar exceed a certain
    length or contain special characters (parentheses, punctuation, etc.).

    Short alphanumeric keys ("1.1", "2.3", …) are entirely safe for grammar generation.
    The mapping from code to full criterion text is stored separately in criterion_map
    (built in submit_to_llm) and injected into the template context for display.

    This approach still enforces at the grammar level that:
      - every criterion is present exactly once (object keys are unique and required),
      - no criterion can be invented (only the declared property names are valid).
    """
    per_criterion_schema = {
        "type": "object",
        "properties": {
            "assessment": {
                "type": "string",
                "enum": ["Strong evidence", "Partial evidence", "Not evident"],
            },
            "commentary": {"type": "string"},
            "confidence": {
                "type": "string",
                "enum": ["high", "medium", "low"],
            },
        },
        "required": ["assessment", "commentary", "confidence"],
    }
    codes = [f"{band_idx}.{crit_idx}" for crit_idx in range(1, len(criteria_list) + 1)]
    return {
        "type": "object",
        "properties": {
            "band_assessment": {
                "type": "string",
                "enum": [
                    "Criteria clearly demonstrated",
                    "Criteria mostly demonstrated",
                    "Some criteria demonstrated",
                    "No evidence",
                ],
            },
            "criteria": {
                "type": "object",
                "properties": {code: per_criterion_schema for code in codes},
                "required": codes,
            },
        },
        "required": ["band_assessment", "criteria"],
    }


# JSON Schema for the llama-server /v1/chat/completions grading call.
# Schema is converted to a GBNF grammar by llama-server for constrained generation.
#
# Grading fields only — metadata (word count, GenAI statement, preface) is extracted
# by a separate dedicated _LLM_METADATA_SCHEMA call and merged into the result dict
# after the grading call completes.
#
# Top-level structure:
#   report_summary   — 1-2 paragraph narrative description of the report content
#   classification   — recommended grade band (enum-constrained to rubric bands)
#   overall_reasoning — one paragraph justifying the recommended band
#   bands            — fixed-key object, one property per rubric band, each containing
#                      a band_assessment summary and a 'criteria' object.
#                      'criteria' is keyed by short numeric codes ("1.1", "1.2", …)
#                      rather than full criterion text — long property-name literals
#                      trigger a SIGSEGV in the llama.cpp GBNF generator.
#   caveats          — free-text limitations or caveats


def _make_llm_response_schema(rubric) -> dict:
    prompt_bands = rubric.to_prompt_bands()
    return {
        "type": "object",
        "properties": {
            "report_summary": {"type": "string"},
            "classification": {
                "type": "string",
                "enum": [band["band"] for band in prompt_bands],
            },
            "overall_reasoning": {"type": "string"},
            "bands": {
                "type": "object",
                "properties": {
                    band["band"]: _make_band_schema(band_idx, band["criteria"])
                    for band_idx, band in enumerate(prompt_bands, start=1)
                },
                "required": [band["band"] for band in prompt_bands],
            },
            "caveats": {"type": "string"},
        },
        "required": [
            "report_summary",
            "classification",
            "overall_reasoning",
            "bands",
            "caveats",
        ],
    }


def _validate_llm_response(data: dict) -> bool:
    """Return True if the LLM JSON response has all required top-level keys."""
    return _REQUIRED_JSON_KEYS.issubset(data.keys())


# ---------------------------------------------------------------------------
# Feedback LLM helpers.
# ---------------------------------------------------------------------------

_LLM_FEEDBACK_RESPONSE_SCHEMA = {
    "type": "object",
    "properties": {
        "positive_feedback": {
            "type": "array",
            "items": {"type": "string"},
            "minItems": 2,
            "maxItems": 3,
        },
        "improvements": {
            "type": "array",
            "items": {"type": "string"},
            "minItems": 2,
            "maxItems": 3,
        },
    },
    "required": ["positive_feedback", "improvements"],
}

_REQUIRED_FEEDBACK_JSON_KEYS = {"positive_feedback", "improvements"}


def _build_feedback_system_prompt(truncated: bool) -> str:
    """Construct the LLM system prompt for the feedback-generation pass."""
    truncation_note = (
        "\n\nNOTE: The document text provided may have been truncated. "
        "If you see the marker '[... middle section omitted due to length ...]', "
        "the middle portion of the document has been removed. Base your feedback "
        "only on the text that is visible."
        if truncated
        else ""
    )

    return f"""You are an experienced academic supervisor reading a student project report. \
Your task is to provide concise, specific formative feedback that will help the student \
understand what they did well and how they could improve.

## Instructions

Produce exactly 2–3 items of **positive feedback** and exactly 2–3 **suggestions for improvement**.

Each item must:
- Be specific to the actual content of this report — name the topic, technique, section, \
or argument you are referring to. Generic phrases such as "good introduction" or \
"consider improving your writing" are not acceptable.
- Be written in plain English, addressed to the student, in one or two sentences.
- For improvements: be actionable — say what the student should do differently, not just \
that something is missing.

## Response format

Your response must be a JSON object with exactly two fields:

- "positive_feedback": an array of 2–3 strings, each being one item of positive feedback.
- "improvements": an array of 2–3 strings, each being one actionable improvement suggestion.{truncation_note}"""


def _build_feedback_user_prompt(document_text: str) -> str:
    return f"Please provide feedback on the following report:\n\n---\n\n{document_text}\n\n---"


def _validate_feedback_response(data: dict) -> bool:
    """Return True if the feedback JSON response has all required top-level keys."""
    return _REQUIRED_FEEDBACK_JSON_KEYS.issubset(data.keys())


# ---------------------------------------------------------------------------
# Shared LLM call helper.
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Feedback deduplication helper.
# ---------------------------------------------------------------------------


def _deduplicate_feedback(items: list[str], max_items: int) -> list[str]:
    """
    Deduplicate feedback strings and return at most *max_items* entries.
    Longer items are preferred as a proxy for specificity.
    """
    seen: set[str] = set()
    unique: list[str] = []
    for item in items:
        normalised = item.strip().lower()
        if normalised not in seen:
            seen.add(normalised)
            unique.append(item.strip())
    unique.sort(key=len, reverse=True)
    return unique[:max_items]


# ---------------------------------------------------------------------------
# Regex patterns for metadata region location.
# ---------------------------------------------------------------------------

_PREFACE_HEADING_RE = re.compile(
    r"^\s*(preface|personal\s+statement"
    r"|declaration(?:\s+of\s+(?:originality|contribution))?"
    r"|acknowledgements?(?:\s+and\s+contribution(?:\s+statement)?)?"
    r"|contribution\s+statement|author.?s?\s+contribution"
    r"|personal\s+contribution)\s*$",
    re.IGNORECASE | re.MULTILINE,
)

_AI_HEADING_RE = re.compile(
    r"^\s*((?:use\s+of\s+)?(?:generative\s+)?ai(?:\s+tools?)?(?:\s+(?:statement|declaration|disclosure|policy|use))?"
    r"|statement\s+on\s+(?:generative\s+)?ai(?:\s+use)?"
    r"|declaration\s+on\s+(?:generative\s+)?ai(?:\s+use)?"
    r"|chatgpt\s+(?:statement|use|policy)"
    r"|large\s+language\s+model\s+(?:statement|use|policy))\s*$",
    re.IGNORECASE | re.MULTILINE,
)

_AI_KEYWORD_SENTENCE_RE = re.compile(
    r"(?:[A-Z][^.!?\n]*)?"
    r"(?:generative\s+ai|chatgpt|gpt-?\d*|large\s+language\s+model"
    r"|github\s+copilot|claude\s+ai|gemini\s+ai|copilot\s+ai)"
    r"[^.!?\n]*[.!?]",
    re.IGNORECASE,
)


# ---------------------------------------------------------------------------
# Metadata region extraction (regex-based location step).
# ---------------------------------------------------------------------------


def _extract_metadata_regions(text: str) -> str:
    """
    Locate candidate text regions for the dedicated metadata LLM call using regex.

    Returns a condensed string containing:
      - front matter up to the first chapter heading (or first 1 500 words),
      - any paragraph block whose heading matches a preface pattern,
      - any paragraph block whose heading matches an AI declaration pattern,
      - fallback: sentences containing AI-related keywords if no heading was found.

    This text is passed verbatim to the metadata LLM; it is never submitted to
    the main assessment LLM.
    """
    paragraphs = [p for p in re.split(r"\n\s*\n", text) if p.strip()]

    # Collect front matter: stop at "Chapter 1" / "Introduction" or after 1 500 words.
    front_end_idx = len(paragraphs)
    word_count = 0
    for i, para in enumerate(paragraphs):
        if re.match(
            r"^\s*(chapter\s+1\b|1[.\s]\s*introduction\b|introduction\s*$)",
            para,
            re.IGNORECASE,
        ):
            front_end_idx = i
            break
        word_count += len(para.split())
        front_end_idx = i + 1
        if word_count >= 1500:
            break

    candidate_indices: set[int] = set(range(front_end_idx))

    # Locate named preface and AI sections anywhere in the document.
    for i, para in enumerate(paragraphs):
        stripped = para.strip()
        if not (_PREFACE_HEADING_RE.match(stripped) or _AI_HEADING_RE.match(stripped)):
            continue
        candidate_indices.add(i)
        for j in range(i + 1, min(i + 25, len(paragraphs))):
            body = paragraphs[j].strip()
            # Stop when we reach what looks like the next section heading:
            # a single short line without terminal punctuation.
            is_next_heading = (
                "\n" not in body
                and len(body.split()) <= 10
                and not body.endswith((".", ",", ";", ":", "?", "!"))
            )
            if is_next_heading and j > i + 1:
                break
            candidate_indices.add(j)

    region_text = "\n\n".join(paragraphs[i] for i in sorted(candidate_indices))

    # If no named AI section was found, append any AI-keyword sentences from the
    # full document as a last-resort fallback.
    ai_section_found = any(
        _AI_HEADING_RE.match(paragraphs[i].strip()) for i in candidate_indices
    )
    if not ai_section_found:
        ai_sentences = _AI_KEYWORD_SENTENCE_RE.findall(text)
        if ai_sentences:
            region_text += (
                "\n\n---\n\nPossible AI use statements found elsewhere in document:\n"
                + "\n".join(s.strip() for s in ai_sentences[:5])
            )

    return region_text


# ---------------------------------------------------------------------------
# Document chunking.
# ---------------------------------------------------------------------------


def _build_chunks(text: str, max_words: int) -> list[str]:
    """
    Split *text* into chunks of at most *max_words* words, breaking at paragraph
    boundaries (double newlines) where possible.  Falls back to sentence
    boundaries for paragraphs that individually exceed *max_words*.
    """
    paragraphs = [p for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[str] = []
    current_paras: list[str] = []
    current_words = 0

    for para in paragraphs:
        w = len(para.split())
        if w > max_words:
            # Oversized single paragraph — split at sentence boundaries.
            if current_paras:
                chunks.append("\n\n".join(current_paras))
                current_paras, current_words = [], 0
            sentences = re.split(r"(?<=[.!?])\s+", para)
            for sent in sentences:
                sw = len(sent.split())
                if current_words + sw > max_words and current_paras:
                    chunks.append("\n\n".join(current_paras))
                    current_paras, current_words = [sent], sw
                else:
                    current_paras.append(sent)
                    current_words += sw
        elif current_words + w > max_words and current_paras:
            chunks.append("\n\n".join(current_paras))
            current_paras, current_words = [para], w
        else:
            current_paras.append(para)
            current_words += w

    if current_paras:
        chunks.append("\n\n".join(current_paras))

    return chunks


# ---------------------------------------------------------------------------
# Map-phase (chunk evidence extraction) helpers.
# ---------------------------------------------------------------------------

_LLM_CHUNK_SCHEMA = {
    "type": "object",
    "properties": {
        "evidence": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "criterion_code": {"type": "string"},
                    "excerpt": {"type": "string", "maxLength": 500},
                    "observation": {"type": "string", "maxLength": 300},
                    "polarity": {
                        "type": "string",
                        "enum": ["supports", "undermines", "neutral"],
                    },
                },
                "required": ["criterion_code", "excerpt", "observation", "polarity"],
            },
        },
    },
    "required": ["evidence"],
}


def _build_chunk_system_prompt(chunk_idx: int, total_chunks: int, rubric) -> str:
    """Compact system prompt for the map-phase evidence extraction call."""
    criterion_lines = []
    for band_idx, band in enumerate(rubric.to_prompt_bands(), start=1):
        for crit_idx, criterion in enumerate(band["criteria"], start=1):
            criterion_lines.append(f"  {band_idx}.{crit_idx}: {criterion}")
    criterion_list = "\n".join(criterion_lines)

    return f"""You are extracting evidence from a portion of a student project report.

This is chunk {chunk_idx + 1} of {total_chunks}. Do NOT assess or classify the work overall.

For each passage that is relevant to any of the criteria listed below, record one evidence entry:
  - criterion_code: the numeric code from the list (e.g. "1.1", "3.4")
  - excerpt: verbatim quote of at most 2 sentences from this chunk
  - observation: one sentence explaining why the passage is relevant
  - polarity: "supports" if the passage shows the criterion is met, "undermines" if it shows the criterion is not met, or "neutral"

If no genuine evidence exists for a criterion in this chunk, do not record an entry for it. An empty evidence array is correct when the chunk contains no relevant passages.

## Criteria

{criterion_list}"""


def _build_chunk_user_prompt(chunk_text: str, chunk_idx: int, total_chunks: int) -> str:
    return (
        f"Please extract evidence from the following text "
        f"(chunk {chunk_idx + 1} of {total_chunks}):\n\n---\n\n{chunk_text}\n\n---"
    )


# ---------------------------------------------------------------------------
# Metadata extraction helpers.
# ---------------------------------------------------------------------------

_LLM_METADATA_SCHEMA = {
    "type": "object",
    "properties": {
        "stated_word_count_found": {"type": "boolean"},
        "stated_word_count": {"type": "integer"},
        "genai_statement_found": {"type": "boolean"},
        "genai_statement": {"type": "string"},
        "preface_found": {"type": "boolean"},
        "preface_precis": {"type": "string"},
    },
    "required": [
        "stated_word_count_found",
        "stated_word_count",
        "genai_statement_found",
        "genai_statement",
        "preface_found",
        "preface_precis",
    ],
}


def _build_metadata_system_prompt() -> str:
    return """You are extracting structured metadata from a student project report.

Your task is to identify and extract the following items, if present in the provided text:

1. **Stated word count**: An explicit word count declared by the student (e.g. under a label such as "Word count:", "Total word count:", or similar). Record the integer value if found.

2. **Generative AI statement**: Any statement about the student's use — or explicit non-use — of generative AI tools such as ChatGPT, GitHub Copilot, or similar. Copy the verbatim text of the statement, or its first sentence if it is very long. This is a personal declaration by the student, not a reference to AI in the domain of study.

3. **Preface / personal contribution statement**: A section (typically headed "Preface", "Personal Statement", "Declaration of Contribution", or similar) describing the student's personal contribution to the project work. Provide a 1–3 sentence précis of what the student claims they personally did.

Record items not found as: stated_word_count_found = false, stated_word_count = 0, genai_statement_found = false, genai_statement = "", preface_found = false, preface_precis = ""."""


def _build_metadata_user_prompt(candidate_text: str) -> str:
    return (
        f"Please extract metadata from the following document excerpt:\n\n"
        f"---\n\n{candidate_text}\n\n---"
    )


# ---------------------------------------------------------------------------
# Map-phase evidence aggregation helpers.
# ---------------------------------------------------------------------------


def _merge_chunk_evidence(chunk_results: dict, all_criterion_codes: frozenset) -> dict:
    """
    Aggregate per-chunk evidence from the map phase.

    Preponderance is preserved: if a criterion is overwhelmingly supported
    (e.g. 6 'supports', 0 'undermines') that dominance is reflected in both
    the retained excerpts and the count header written into the synthesis prompt.
    The retention policy is:
      - Keep at most _MAX_EVIDENCE_PER_CRITERION entries total per criterion.
      - Always include at least one entry from every polarity that exists
        (genuine minority evidence is informative).
      - Fill remaining slots with entries from the dominant polarity.

    Metadata is merged with OR semantics for booleans and first-found semantics
    for string values (the first chunk that identifies a field wins).

    Returns:
        {
          "criteria": {
            code: {
              "entries": [evidence_dict, ...],
              "counts": {"supports": N, "undermines": N, "neutral": N,
                         "chunks_with_evidence": N}
            }
          },
          "metadata": { stated_word_count_found, stated_word_count,
                         genai_statement_found, genai_statement,
                         preface_found, preface_precis }
        }
    """
    from collections import defaultdict

    raw: dict[str, dict[str, list]] = defaultdict(
        lambda: {"supports": [], "undermines": [], "neutral": []}
    )
    chunks_with_evidence: dict[str, set] = defaultdict(set)
    metadata: dict = {
        "stated_word_count_found": False,
        "stated_word_count": None,
        "genai_statement_found": False,
        "genai_statement": "",
        "preface_found": False,
        "preface_precis": "",
    }

    for chunk_idx_str, chunk_data in chunk_results.items():
        chunk_idx = int(chunk_idx_str)

        # Merge metadata (OR / first-found).
        hits = chunk_data.get("metadata_hits", {})
        if (
            hits.get("stated_word_count_found")
            and not metadata["stated_word_count_found"]
        ):
            metadata["stated_word_count_found"] = True
            metadata["stated_word_count"] = hits.get("stated_word_count")
        if hits.get("genai_statement_found") and not metadata["genai_statement_found"]:
            metadata["genai_statement_found"] = True
            metadata["genai_statement"] = hits.get("genai_statement", "")
        if hits.get("preface_found") and not metadata["preface_found"]:
            metadata["preface_found"] = True
            metadata["preface_precis"] = hits.get("preface_precis", "")

        # Accumulate evidence.
        for entry in chunk_data.get("evidence", []):
            code = entry.get("criterion_code", "")
            polarity = entry.get("polarity", "neutral")
            if code not in all_criterion_codes:
                continue
            if polarity not in ("supports", "undermines", "neutral"):
                polarity = "neutral"
            raw[code][polarity].append(entry)
            chunks_with_evidence[code].add(chunk_idx)

    # Apply retention policy.
    criteria: dict = {}
    for code, polarities in raw.items():
        supports = polarities["supports"]
        undermines = polarities["undermines"]
        neutral = polarities["neutral"]

        counts = {
            "supports": len(supports),
            "undermines": len(undermines),
            "neutral": len(neutral),
            "chunks_with_evidence": len(chunks_with_evidence[code]),
        }

        # Sort polarities by count descending to find dominant.
        by_count = sorted(
            [("supports", supports), ("undermines", undermines), ("neutral", neutral)],
            key=lambda x: len(x[1]),
            reverse=True,
        )
        dominant_name, dominant_entries = by_count[0]

        # One entry from each minority polarity that actually exists.
        minority: list[dict] = []
        for _, entries in by_count[1:]:
            if entries:
                minority.append(entries[0])

        # Fill remaining slots with dominant entries.
        remaining = max(_MAX_EVIDENCE_PER_CRITERION - len(minority), 0)
        retained = dominant_entries[:remaining] + minority

        criteria[code] = {"entries": retained, "counts": counts}

    return {"criteria": criteria, "metadata": metadata}


# ---------------------------------------------------------------------------
# Synthesis (reduce-phase) helpers.
# ---------------------------------------------------------------------------


def _build_synthesis_evidence_text(merged: dict, total_chunks: int, rubric) -> str:
    """
    Format aggregated chunk evidence as structured text for the synthesis prompt.
    This replaces the document text in the synthesis user prompt.

    Per-criterion count headers convey preponderance to the synthesis LLM so it
    can distinguish criteria with overwhelming evidence from those that are
    genuinely ambiguous.
    """
    criteria_data: dict = merged["criteria"]
    metadata: dict = merged["metadata"]
    chunk_word = "chunk" if total_chunks == 1 else "chunks"

    lines = [f"## Extracted evidence ({total_chunks} {chunk_word})", ""]

    for band_idx, band in enumerate(rubric.to_prompt_bands(), start=1):
        lines.append(f"### Band {band_idx} — {band['band']}")
        for crit_idx, criterion in enumerate(band["criteria"], start=1):
            code = f"{band_idx}.{crit_idx}"
            crit_data = criteria_data.get(code)
            lines.append(f"\n**{code}** {criterion}")
            if crit_data is None:
                lines.append("  (No evidence found in any chunk)")
                continue

            counts = crit_data["counts"]
            count_summary = (
                f"{counts['supports']} supporting, "
                f"{counts['undermines']} undermining, "
                f"{counts['neutral']} neutral "
                f"across {counts['chunks_with_evidence']} of {total_chunks} {chunk_word}"
            )
            lines.append(f"  Evidence count: {count_summary}")

            for entry in crit_data["entries"]:
                tag = entry["polarity"].upper()
                excerpt = entry["excerpt"]
                if len(excerpt) > _MAX_EXCERPT_CHARS:
                    excerpt = excerpt[:_MAX_EXCERPT_CHARS] + "…"
                lines.append(f'  [{tag}] "{excerpt}"')
                lines.append(f"  {entry['observation']}")
        lines.append("")

    lines.append("## Document metadata")
    if metadata["stated_word_count_found"]:
        lines.append(f"- Stated word count: {metadata['stated_word_count']}")
    else:
        lines.append("- Stated word count: not found")

    if metadata["genai_statement_found"]:
        lines.append(f'- GenAI statement: "{metadata["genai_statement"]}"')
    else:
        lines.append("- GenAI statement: not found")

    if metadata["preface_found"]:
        lines.append(f'- Preface/personal contribution: "{metadata["preface_precis"]}"')
    else:
        lines.append("- Preface/personal contribution: not found")

    return "\n".join(lines)


def _build_synthesis_user_prompt(evidence_text: str) -> str:
    return (
        "The full report has been pre-read and relevant passages extracted by assessment criterion.\n"
        "Use the evidence below to produce your final grade band assessment.\n\n"
        "The evidence entries include counts of supporting and undermining passages per criterion. "
        "Use these counts to judge preponderance, not just the individual quoted excerpts: "
        "a criterion with 6 supporting and 0 undermining passages deserves 'Strong evidence', "
        "not 'Partial evidence'.\n\n"
        "For criteria with no evidence entries, use 'Not evident' with low confidence.\n\n"
        f"{evidence_text}"
    )


# ---------------------------------------------------------------------------
# Celery task registration.
# ---------------------------------------------------------------------------


def _dispatch_process_report(record_id: int) -> None:
    """Fire-and-forget dispatch of process_report.process after language analysis completes."""
    try:
        celery = current_app.extensions["celery"]
        task = celery.tasks["app.tasks.process_report.process"]
        task.apply_async(args=[record_id])
    except Exception as exc:
        current_app.logger.exception(
            f"Failed to dispatch process_report for record_id={record_id}", exc_info=exc
        )


def register_language_analysis_tasks(celery):

    @celery.task(bind=True, default_retry_delay=30)
    def download_and_extract(self, record_id: int):
        """
        Stage 1 (llm_tasks queue): download the submitted report asset and
        extract plain text from it.  The extracted text is stored in the
        MongoDB scraped-text cache for use by subsequent pipeline stages.
        """
        self.update_state(
            state=states.STARTED, meta={"msg": "Downloading report asset"}
        )

        _r = None
        try:
            _r = get_pipeline_redis()
        except Exception:
            pass
        _t0 = record_step_start(_r, record_id, "download_and_extract")

        try:
            record: SubmissionRecord = (
                db.session.query(SubmissionRecord).filter_by(id=record_id).first()
            )
        except SQLAlchemyError as exc:
            record_step_end(_r, record_id, "download_and_extract", _t0, error=repr(exc))
            current_app.logger.exception(
                "SQLAlchemyError in download_and_extract", exc_info=exc
            )
            raise self.retry()

        if record is None:
            record_step_end(
                _r,
                record_id,
                "download_and_extract",
                _t0,
                error="SubmissionRecord not found",
            )
            raise Exception(
                f"language_analysis.download_and_extract: SubmissionRecord #{record_id} not found"
            )

        if record.report is None:
            record_step_end(
                _r,
                record_id,
                "download_and_extract",
                _t0,
                error="SubmissionRecord has no report",
            )
            raise Exception(
                f"language_analysis.download_and_extract: SubmissionRecord #{record_id} has no report"
            )

        asset = record.report

        # Idempotency check: skip download if MongoDB already has current scraped
        # text for this exact asset.  The uploaded asset is immutable, so a matching
        # asset_id means the cached text is still valid.
        cached = get_scraped_text(record_id)
        if cached is not None and cached.get("asset_id") == asset.id:
            current_app.logger.info(
                f"language_analysis.download_and_extract: record #{record_id} — "
                f"scraped text cache is current (asset_id={asset.id}); skipping download"
            )
            record_step_end(_r, record_id, "download_and_extract", _t0)
            return

        storage = current_app.config["OBJECT_STORAGE_ASSETS"]
        adapter = AssetCloudAdapter(
            asset,
            storage,
            audit_data=f"language_analysis.download_and_extract (record #{record_id})",
        )

        if not adapter.exists():
            raise Exception(
                f"language_analysis.download_and_extract: report asset not found in object store for record #{record_id}"
            )

        raw_text = ""
        page_count = 0
        errors = []

        mimetype = (asset.mimetype or "").lower()
        _t_extraction = time.monotonic()

        try:
            with adapter.download_to_scratch() as scratch:
                path = str(scratch.path)
                if "pdf" in mimetype or path.lower().endswith(".pdf"):
                    raw_text, page_count = _extract_pdf_text(path)
                elif (
                    "word" in mimetype
                    or "officedocument" in mimetype
                    or path.lower().endswith((".docx", ".doc"))
                ):
                    raw_text, page_count = _extract_docx_text(path)
                else:
                    # Fall back to PDF extraction and log a warning
                    current_app.logger.warning(
                        f"language_analysis: unknown mimetype '{mimetype}' for record #{record_id}; attempting PDF extraction"
                    )
                    try:
                        raw_text, page_count = _extract_pdf_text(path)
                    except Exception as exc2:
                        errors.append(
                            {
                                "stage": "extract",
                                "type": type(exc2).__name__,
                                "message": str(exc2),
                            }
                        )
        except Exception as exc:
            errors.append(
                {"stage": "download", "type": type(exc).__name__, "message": str(exc)}
            )

        # Cache extracted text in MongoDB for use by subsequent pipeline stages
        # and future pairwise similarity analysis.
        normalized_text = unicodedata.normalize("NFKC", raw_text)
        store_scraped_text(record_id, asset.id, mimetype, normalized_text, page_count)

        data = record.language_analysis_data
        data["_page_count"] = page_count
        if errors:
            data.setdefault("errors", []).extend(errors)
        data.setdefault("timings", {})["extraction_s"] = round(
            time.monotonic() - _t_extraction, 1
        )
        record.set_language_analysis_data(data)

        try:
            db.session.commit()
        except SQLAlchemyError as exc:
            db.session.rollback()
            record_step_end(_r, record_id, "download_and_extract", _t0, error=repr(exc))
            current_app.logger.exception(
                "SQLAlchemyError committing extracted text", exc_info=exc
            )
            raise self.retry()

        record_step_end(_r, record_id, "download_and_extract", _t0)

    # -----------------------------------------------------------------------

    @celery.task(bind=True, default_retry_delay=30)
    def compute_statistics(self, record_id: int):
        """
        Stage 2 (default queue): compute all statistical language metrics from
        the extracted text stored in the JSON blob.  Errors in individual
        computation steps are caught and recorded; the task does not abort.
        """
        self.update_state(
            state=states.STARTED, meta={"msg": "Computing language statistics"}
        )

        _r = None
        try:
            _r = get_pipeline_redis()
        except Exception:
            pass
        _t0 = record_step_start(_r, record_id, "compute_statistics")

        try:
            record: SubmissionRecord = (
                db.session.query(SubmissionRecord).filter_by(id=record_id).first()
            )
        except SQLAlchemyError as exc:
            record_step_end(_r, record_id, "compute_statistics", _t0, error=repr(exc))
            current_app.logger.exception(
                "SQLAlchemyError in compute_statistics", exc_info=exc
            )
            raise self.retry()

        if record is None:
            record_step_end(
                _r,
                record_id,
                "compute_statistics",
                _t0,
                error="SubmissionRecord not found",
            )
            raise Exception(
                f"language_analysis.compute_statistics: SubmissionRecord #{record_id} not found"
            )

        # Idempotency check: skip if statistics are already present at the current version
        if (
            record.stats_present
            and record.stats_algorithm_version == STATS_ALGORITHM_VERSION
        ):
            current_app.logger.info(
                f"language_analysis.compute_statistics: skipping record #{record_id} — "
                f"stats already present at algorithm_version={STATS_ALGORITHM_VERSION}"
            )
            record_step_end(_r, record_id, "compute_statistics", _t0)
            return

        data = record.language_analysis_data
        _cached = get_scraped_text(record_id)
        raw_text: str = _cached["scraped_text"] if _cached else ""
        errors: list = data.get("errors", [])

        metrics: dict = {}
        references_info: dict = {}
        patterns_info: dict = {}

        # --- split into core, references, and appendices ---------------------
        # _split_document() implements the priority-ordered 3-way split:
        #   _core       — main body text (no references, no appendices)
        #   _references — bibliography / reference list
        #   _appendices — appendix sections (may be empty)
        _core, _references, _appendices = _split_document(raw_text)

        # Strip math-extraction noise from the core body text.  Appendix text
        # is also stripped separately for appendix word counting.
        # See _strip_math_lines() for full rationale and threshold choice.
        clean_core_text = _strip_math_lines(_core)

        # Content text for figure/table detection: core + appendices (not refs).
        _content_text = _core + ("\n\n" + _appendices if _appendices else "")

        # Stripped content text for lexical diversity and sentence-structure metrics.
        # Appendices are included because they are the student's own writing, consistent
        # with the text submitted to the LLM.  Word count stays core-only (see above).
        clean_content_text = _strip_math_lines(_content_text)

        _t_counting = time.monotonic()

        # --- word count (core body only — appendices excluded) ---------------
        try:
            wc = _word_count(clean_core_text)
            metrics["word_count"] = wc
            # Appendix word count stored separately so the UI can surface both.
            if _appendices:
                appendix_wc = _word_count(_strip_math_lines(_appendices))
                if appendix_wc > 0:
                    metrics["appendix_word_count"] = appendix_wc
        except Exception as exc:
            errors.append(
                {"stage": "word_count", "type": type(exc).__name__, "message": str(exc)}
            )
            metrics["word_count"] = None

        # --- bibliography count and citation check ----------------------------
        try:
            ref_count, ref_keys = _count_bibliography(_references)
            metrics["reference_count"] = ref_count

            uncited = _check_uncited(_core, ref_keys)
            references_info["uncited"] = uncited
        except Exception as exc:
            errors.append(
                {"stage": "references", "type": type(exc).__name__, "message": str(exc)}
            )
            metrics["reference_count"] = None
            references_info["uncited"] = []

        # --- figure and table cross-reference check --------------------------
        try:
            uncaptioned_figs, uncaptioned_tabs = _check_figure_table_refs(_content_text)
            # Count distinct canonical labels (e.g. "3.1", "A.1") in core + appendices.
            fig_labels = {m.group(1) for m in _FIG_REF.finditer(_content_text)}
            tab_labels = {m.group(1) for m in _TAB_REF.finditer(_content_text)}
            metrics["figure_count"] = len(fig_labels)
            metrics["table_count"] = len(tab_labels)
            references_info["uncaptioned_figures"] = uncaptioned_figs
            references_info["uncaptioned_tables"] = uncaptioned_tabs
        except Exception as exc:
            errors.append(
                {
                    "stage": "figure_table_refs",
                    "type": type(exc).__name__,
                    "message": str(exc),
                }
            )
            metrics["figure_count"] = None
            metrics["table_count"] = None
            references_info["uncaptioned_figures"] = []
            references_info["uncaptioned_tables"] = []

        # counting_s covers: word count, bibliography, figure/table refs, pattern
        # matching — all fast regex/counting operations with no NLP model load.
        _t_ai = time.monotonic()

        # --- MATTR and MTLD --------------------------------------------------
        try:
            mattr, mtld = _compute_mattr_mtld(clean_content_text)
            metrics["mattr"] = mattr
            metrics["mtld"] = mtld
        except Exception as exc:
            errors.append(
                {"stage": "mattr_mtld", "type": type(exc).__name__, "message": str(exc)}
            )
            metrics["mattr"] = None
            metrics["mtld"] = None

        # --- burstiness ------------------------------------------------------
        try:
            burstiness_groups, burstiness_aggregate = _compute_burstiness(raw_text)
            metrics["burstiness"] = burstiness_aggregate
            metrics["burstiness_by_group"] = burstiness_groups
        except Exception as exc:
            errors.append(
                {"stage": "burstiness", "type": type(exc).__name__, "message": str(exc)}
            )
            metrics["burstiness"] = None
            metrics["burstiness_by_group"] = {}

        # --- sentence CV -----------------------------------------------------
        try:
            metrics["sentence_cv"] = _compute_sentence_cv(clean_content_text)
        except Exception as exc:
            errors.append(
                {
                    "stage": "sentence_cv",
                    "type": type(exc).__name__,
                    "message": str(exc),
                }
            )
            metrics["sentence_cv"] = None

        # --- pattern matching ------------------------------------------------
        try:
            patterns_info = _count_patterns(raw_text)
        except Exception as exc:
            errors.append(
                {"stage": "patterns", "type": type(exc).__name__, "message": str(exc)}
            )

        # --- classification flags --------------------------------------------
        mattr_flag = classify_mattr(metrics.get("mattr"))
        mtld_flag = classify_mtld(metrics.get("mtld"))
        burst_flag = classify_burstiness(metrics.get("burstiness"))
        cv_flag = classify_sentence_cv(metrics.get("sentence_cv"))

        # Fetch calibrations from the tenant associated with this project class.
        # NLL metrics are not available at this stage; full calibrations will be
        # re-evaluated by submit_to_llm once NLL has been computed.
        calibrations: list = []
        try:
            pclass = record.period.config.project_class
            tenant = pclass.tenant if pclass else None
            calibrations = list(tenant.ai_calibrations) if tenant else []
        except Exception:
            calibrations = []

        ai_result = _ai_concern_flag(
            metrics.get("mattr"),
            metrics.get("mtld"),
            metrics.get("sentence_cv"),
            calibrations,
        )

        flags = {
            "mattr_flag": mattr_flag,
            "mtld_flag": mtld_flag,
            "burstiness_flag": burst_flag,
            "sentence_cv_flag": cv_flag,
            "ai_concern": ai_result["concern"],
            "mahalanobis_sigma": ai_result["sigma"],
            "mahalanobis_pvalue": ai_result["p_value"],
            "calibration_results": ai_result.get("calibration_results", []),
            "bonferroni_k": ai_result.get("bonferroni_k", 0),
            "bonferroni_alpha_medium": ai_result.get("bonferroni_alpha_medium"),
            "bonferroni_alpha_high": ai_result.get("bonferroni_alpha_high"),
        }

        # --- persist ---------------------------------------------------------
        timings = data.get("timings", {})
        timings["counting_s"] = round(_t_ai - _t_counting, 1)
        timings["ai_metrics_s"] = round(time.monotonic() - _t_ai, 1)
        data["timings"] = timings
        data["metrics"] = metrics
        data["flags"] = flags
        data["references"] = references_info
        data["patterns"] = patterns_info
        data["errors"] = errors
        record.set_language_analysis_data(data)
        record.stats_present = True
        record.stats_algorithm_version = STATS_ALGORITHM_VERSION

        try:
            db.session.commit()
        except SQLAlchemyError as exc:
            db.session.rollback()
            record_step_end(_r, record_id, "compute_statistics", _t0, error=repr(exc))
            current_app.logger.exception(
                "SQLAlchemyError committing statistics", exc_info=exc
            )
            raise self.retry()

        record_step_end(_r, record_id, "compute_statistics", _t0)

    # -----------------------------------------------------------------------

    @celery.task(
        bind=True, default_retry_delay=30, soft_time_limit=7200, time_limit=7260
    )
    def submit_to_llm(self, record_id: int):
        """
        Stage 3 (llm_tasks queue): submit the extracted report text to the LLM
        for grade-band assessment.

        Short documents that fit within the context window are submitted in a
        single pass (unchanged behaviour).  Longer documents use a two-phase
        map-reduce strategy:

          Map phase — each chunk is submitted with _LLM_CHUNK_SCHEMA to
                      extract per-criterion evidence fragments.  A dedicated
                      metadata call (regex location + _LLM_METADATA_SCHEMA)
                      extracts the GenAI statement and preface text from the
                      identified front-matter regions.  All intermediate
                      results are persisted to the JSON blob after each chunk
                      so the task is safely resumable on Celery retry.

          Reduce phase — _merge_chunk_evidence() aggregates the map results,
                         preserving preponderance (evidence counts per criterion
                         are passed to the synthesis LLM alongside excerpts).
                         The synthesis call uses the full grading response schema
                         and produces output structurally identical to the
                         single-pass result.

        Retries up to _LLM_RETRY_ATTEMPTS times per LLM call; permanent
        failures set llm_analysis_failed.  Records with llm_analysis_failed=True
        are not retried until an administrator explicitly clears the flag.
        """
        self.update_state(state=states.STARTED, meta={"msg": "Submitting to LLM"})

        _r = None
        try:
            _r = get_pipeline_redis()
        except Exception:
            pass
        _t0 = record_step_start(_r, record_id, "submit_to_llm")

        try:
            record: SubmissionRecord = (
                db.session.query(SubmissionRecord).filter_by(id=record_id).first()
            )
        except SQLAlchemyError as exc:
            current_app.logger.exception(
                "SQLAlchemyError in submit_to_llm", exc_info=exc
            )
            raise self.retry()

        if record is None:
            raise Exception(
                f"language_analysis.submit_to_llm: SubmissionRecord #{record_id} not found"
            )

        # Do not re-attempt if a human administrator has not yet cleared the failure flag
        if record.llm_analysis_failed:
            current_app.logger.info(
                f"language_analysis.submit_to_llm: skipping record #{record_id} — llm_analysis_failed is set"
            )
            record_step_end(_r, record_id, "submit_to_llm", _t0)
            return

        data = record.language_analysis_data

        # Idempotency check: skip if a grading result is already present at the current
        # prompt version.  Bumping PROMPT_VERSION in code is the mechanism to force
        # re-grading across all records.
        if record.llm_grading_present and record.llm_prompt_version == PROMPT_VERSION:
            current_app.logger.info(
                f"language_analysis.submit_to_llm: skipping record #{record_id} — "
                f"grading result already present at prompt_version={PROMPT_VERSION}"
            )
            record_step_end(_r, record_id, "submit_to_llm", _t0)
            return

        _cached = get_scraped_text(record_id)
        raw_text: str = _cached["scraped_text"] if _cached else ""

        # Load rubric while session is still open; snapshot before session.close().
        _rubric_orm = record.period.config.grading_rubric
        rubric_snap: _RubricSnapshot | None = (
            _RubricSnapshot(_rubric_orm) if _rubric_orm is not None else None
        )

        # Build the text for grade-band assessment: core body + appendices.
        # The reference list is excluded (bibliographic entries are noise for
        # the LLM assessor).  Math-extraction artefacts are stripped so that
        # equation fragments do not waste context tokens.
        _core, _references, _appendices = _split_document(raw_text)
        clean_text = _strip_toc_lines(_strip_math_lines(_core))
        if _appendices:
            clean_text = (
                clean_text + "\n\n" + _strip_toc_lines(_strip_math_lines(_appendices))
            )

        context_size: int = current_app.config.get("OLLAMA_CONTEXT_SIZE", 18432)
        base_url: str = current_app.config.get(
            "OLLAMA_BASE_URL", "http://localhost:11434"
        )
        model: str = current_app.config.get("OLLAMA_MODEL", "llama3.1:70b")

        # Release the DB connection back to the pool before the long-running LLM call.
        # The connection can sit idle for 5-10+ minutes during LLM inference, which
        # may exceed the MySQL server-side wait_timeout and leave it stale.  All
        # required data has already been read into local Python variables above.
        # The record will be reloaded from the DB before each subsequent write.
        db.session.close()

        errors: list = data.get("errors", [])

        mean_nll, nll_cv = None, None

        _t_llm = time.monotonic()

        # ----------------------------------------------------------------
        # No-rubric path: run metadata extraction only, skip grading.
        # ----------------------------------------------------------------
        if rubric_snap is None:
            candidate_text = _extract_metadata_regions(clean_text)
            meta_parsed, _, meta_exc, _, _ = _call_llm(
                base_url,
                model,
                _build_metadata_system_prompt(),
                _build_metadata_user_prompt(candidate_text),
                _LLM_METADATA_SCHEMA,
                options={"num_ctx": context_size},
                label=f"submit_to_llm/metadata-only (record #{record_id})",
            )
            data["grading_skipped"] = True
            if meta_parsed is not None:
                data["_llm_metadata"] = meta_parsed
            else:
                current_app.logger.warning(
                    f"language_analysis.submit_to_llm: metadata-only extraction failed "
                    f"for record #{record_id}: {meta_exc}"
                )

            # Store NLL metrics and re-evaluate AI concern with NLL features.
            metrics = data.get("metrics", {})
            metrics["mean_nll"] = mean_nll
            metrics["nll_cv"] = nll_cv
            data["metrics"] = metrics
            data["llm_meta"] = {
                "model": model,
                "context_size": context_size,
                "num_chunks": 0,
            }
            data.setdefault("timings", {})["llm_s"] = round(
                time.monotonic() - _t_llm, 1
            )

            try:
                _pclass = (
                    record.period.config.project_class
                    if record.period and record.period.config
                    else None
                )
                _tenant = _pclass.tenant if _pclass else None
                _calibrations = list(_tenant.ai_calibrations) if _tenant else []
            except Exception:
                _calibrations = []

            _ai_result = _ai_concern_flag(
                metrics.get("mattr"),
                metrics.get("mtld"),
                metrics.get("sentence_cv"),
                _calibrations,
                mean_nll=mean_nll,
                nll_cv=nll_cv,
                llm_model_name=model,
                llm_context_window=context_size,
            )
            _flags = data.get("flags", {})
            _flags["ai_concern"] = _ai_result["concern"]
            _flags["mahalanobis_sigma"] = _ai_result["sigma"]
            _flags["mahalanobis_pvalue"] = _ai_result["p_value"]
            _flags["calibration_results"] = _ai_result.get("calibration_results", [])
            _flags["bonferroni_k"] = _ai_result.get("bonferroni_k", 0)
            _flags["bonferroni_alpha_medium"] = _ai_result.get(
                "bonferroni_alpha_medium"
            )
            _flags["bonferroni_alpha_high"] = _ai_result.get("bonferroni_alpha_high")
            data["flags"] = _flags

            data["errors"] = errors
            data["prompt_version"] = PROMPT_VERSION
            # Clear any stale feedback from a previous run; feedback must re-run
            # against the fresh grading result.
            data.pop("llm_feedback", None)
            record = db.session.get(SubmissionRecord, record_id)
            if record is None:
                raise Exception(
                    f"submit_to_llm: SubmissionRecord #{record_id} not found on reload (no-rubric)"
                )
            record.llm_model_name = model
            record.llm_context_size = context_size
            record.llm_grading_present = True
            record.llm_prompt_version = PROMPT_VERSION
            record.llm_feedback_present = False
            record.llm_feedback_prompt_version = None
            record.set_language_analysis_data(data)
            try:
                db.session.commit()
            except SQLAlchemyError as exc:
                db.session.rollback()
                current_app.logger.exception(
                    "SQLAlchemyError committing no-rubric metadata result", exc_info=exc
                )
                raise self.retry()
            record_step_end(_r, record_id, "submit_to_llm", _t0)
            return

        # ----------------------------------------------------------------
        # Rubric-present grading path.
        # ----------------------------------------------------------------
        all_criterion_codes: frozenset = frozenset(
            f"{band_idx}.{crit_idx}"
            for band_idx, band in enumerate(rubric_snap._bands, start=1)
            for crit_idx in range(1, len(band["criteria"]) + 1)
        )
        n_criteria = sum(len(band["criteria"]) for band in rubric_snap._bands)
        llm_response_schema = _make_llm_response_schema(rubric_snap)

        _grading_prompt = _build_system_prompt(False, rubric_snap)
        _grading_prompt_tokens = int(len(_grading_prompt.split()) * _TOKENS_PER_WORD)
        # Per criterion: ~130 tokens (assessment enum + ~80-word commentary + confidence enum).
        # Fixed overhead: ~700 tokens (summary, classification, overall_reasoning, caveats, JSON framing).
        _single_pass_response_tokens = max(2200, 700 + n_criteria * 130)
        _single_pass_overhead = _grading_prompt_tokens + _single_pass_response_tokens
        single_pass_word_budget = max(
            int(
                (context_size - _single_pass_overhead) / _TOKENS_PER_WORD_CONTENT * 0.85
            ),
            0,
        )
        doc_words = len(clean_text.split())

        accumulated = ""
        last_exc: Exception | None = None
        parsed_result: dict | None = None
        prompt_hash_val: str | None = None
        num_chunks = 1  # updated to actual chunk count on the chunked path
        est_tok: int = 0
        # Token instrumentation — computed on each path and passed to record_step_end.
        _peak_prompt_tokens: int | None = None
        _peak_context_pressure: float | None = None
        _total_est_tokens: int | None = None
        _total_actual_prompt_tokens: int | None = None
        _peak_completion_tokens: int | None = None
        _total_completion_tokens: int | None = None

        if doc_words <= single_pass_word_budget:
            # ----------------------------------------------------------------
            # Single-pass path: document fits within the context window.
            # ----------------------------------------------------------------
            # Step 1: dedicated metadata extraction (same as chunked path step 1).
            candidate_text = _extract_metadata_regions(clean_text)
            metadata_result: dict | None = None
            meta_parsed, _, meta_exc, _, _ = _call_llm(
                base_url,
                model,
                _build_metadata_system_prompt(),
                _build_metadata_user_prompt(candidate_text),
                _LLM_METADATA_SCHEMA,
                options={"num_ctx": context_size},
                label=f"submit_to_llm/metadata (record #{record_id})",
            )
            if meta_parsed is not None:
                metadata_result = meta_parsed
            else:
                current_app.logger.warning(
                    f"language_analysis.submit_to_llm: metadata extraction failed "
                    f"for record #{record_id}: {meta_exc}"
                )

            # Step 2: grading call (grading fields only).
            document_text, was_truncated = _truncate_text(clean_text)
            _system_prompt = _build_system_prompt(was_truncated, rubric_snap)
            prompt_hash_val = _prompt_hash(_system_prompt)
            parsed_result, accumulated, last_exc, est_tok, _sp_actual_usage = _call_llm(
                base_url,
                model,
                _system_prompt,
                _build_user_prompt(document_text),
                llm_response_schema,
                options={"num_ctx": context_size},
                validate_fn=_validate_llm_response,
                label=f"submit_to_llm/single-pass (record #{record_id})",
                user_tokens_per_word=_TOKENS_PER_WORD_CONTENT,
            )

            # Merge metadata into the grading result so downstream code and
            # templates can access all fields from a single dict.
            if parsed_result is not None and metadata_result is not None:
                for field in (
                    "stated_word_count_found",
                    "stated_word_count",
                    "genai_statement_found",
                    "genai_statement",
                    "preface_found",
                    "preface_precis",
                ):
                    parsed_result[field] = metadata_result.get(field)

            # Token instrumentation for single-pass path.
            _total_est_tokens = est_tok
            if _sp_actual_usage is not None:
                _pt = _sp_actual_usage.get("prompt_tokens")
                if _pt is not None:
                    _peak_prompt_tokens = _pt
                    _total_actual_prompt_tokens = _pt
                    _peak_context_pressure = _pt / context_size
                _ct = _sp_actual_usage.get("completion_tokens")
                if _ct is not None:
                    _peak_completion_tokens = _ct
                    _total_completion_tokens = _ct

        else:
            # ----------------------------------------------------------------
            # Chunked map-reduce path: document exceeds single-pass budget.
            # ----------------------------------------------------------------
            _sample_chunk_prompt = _build_chunk_system_prompt(0, 1, rubric_snap)
            _chunk_prompt_tokens = int(
                len(_sample_chunk_prompt.split()) * _TOKENS_PER_WORD
            )
            # Per criterion: up to _MAX_EVIDENCE_PER_CRITERION entries.  Each entry carries a
            # 2-sentence verbatim excerpt (~80 tokens) + observation (~30 tokens) + overhead,
            # so ~220 tokens/criterion at the average 2-entry density is more realistic than the
            # old 150.  The higher floor (1200) covers small rubrics safely.
            _map_response_tokens = max(1200, 500 + n_criteria * 220)
            _map_overhead = _chunk_prompt_tokens + _map_response_tokens
            chunk_word_budget = max(
                int((context_size - _map_overhead) / _TOKENS_PER_WORD_CONTENT * 0.85),
                500,
            )
            chunks = _build_chunks(clean_text, chunk_word_budget)
            total_chunks = len(chunks)
            num_chunks = total_chunks
            current_app.logger.info(
                f"language_analysis.submit_to_llm: record #{record_id} — "
                f"{doc_words} words, {total_chunks} chunk(s) of ~{chunk_word_budget} words "
                f"(context_size={context_size})"
            )

            # -- Step 1: dedicated metadata extraction (regex → LLM) ------
            metadata_result: dict | None = data.get("_llm_metadata")
            if metadata_result is None:
                candidate_text = _extract_metadata_regions(clean_text)
                meta_parsed, _, meta_exc, _, _ = _call_llm(
                    base_url,
                    model,
                    _build_metadata_system_prompt(),
                    _build_metadata_user_prompt(candidate_text),
                    _LLM_METADATA_SCHEMA,
                    options={"num_ctx": context_size},
                    label=f"submit_to_llm/metadata (record #{record_id})",
                )
                if meta_parsed is not None:
                    metadata_result = meta_parsed
                else:
                    # Non-fatal: map phase metadata_hits act as a fallback.
                    current_app.logger.warning(
                        f"language_analysis.submit_to_llm: metadata extraction failed "
                        f"for record #{record_id}: {meta_exc}; "
                        f"map-phase metadata_hits will be used instead"
                    )
                    metadata_result = {
                        "stated_word_count_found": False,
                        "stated_word_count": None,
                        "genai_statement_found": False,
                        "genai_statement": "",
                        "preface_found": False,
                        "preface_precis": "",
                    }
                data["_llm_metadata"] = metadata_result
                record = db.session.get(SubmissionRecord, record_id)
                if record is None:
                    raise Exception(
                        f"submit_to_llm: SubmissionRecord #{record_id} not found on reload (metadata)"
                    )
                record.set_language_analysis_data(data)
                try:
                    db.session.commit()
                except SQLAlchemyError as exc:
                    db.session.rollback()
                    current_app.logger.exception(
                        "SQLAlchemyError committing metadata result", exc_info=exc
                    )
                    raise self.retry()
                db.session.close()

            # -- Step 2: map phase (per-chunk evidence extraction) ---------
            chunk_state = data.get("_llm_chunks", {})

            # Reset persisted state if chunk topology changed between retries
            # (e.g. OLLAMA_CONTEXT_SIZE was adjusted by an administrator).
            if chunk_state.get("total_chunks") != total_chunks:
                chunk_state = {
                    "total_chunks": total_chunks,
                    "chunk_word_budget": chunk_word_budget,
                    "completed": [],
                    "results": {},
                }

            completed_chunks: set[int] = set(chunk_state.get("completed", []))
            chunk_results: dict = chunk_state.get("results", {})
            chunk_failed = False
            chunk_failure_reason = ""
            _chunk_est_tokens: list[int] = []
            _chunk_actual_tokens: list[int | None] = []
            _chunk_completion_tokens: list[int | None] = []

            for idx, chunk_text in enumerate(chunks):
                if idx in completed_chunks:
                    continue  # already persisted on a previous Celery attempt

                chunk_parsed, accumulated, last_exc, est_tok, _chunk_actual_usage = _call_llm(
                    base_url,
                    model,
                    _build_chunk_system_prompt(idx, total_chunks, rubric_snap),
                    _build_chunk_user_prompt(chunk_text, idx, total_chunks),
                    _LLM_CHUNK_SCHEMA,
                    options={"num_ctx": context_size},
                    label=f"submit_to_llm/chunk {idx + 1}/{total_chunks} (record #{record_id})",
                    user_tokens_per_word=_TOKENS_PER_WORD_CONTENT,
                )

                if chunk_parsed is None:
                    chunk_failed = True
                    chunk_failure_reason = f"chunk {idx + 1}/{total_chunks} failed (~{est_tok} est. input tokens): {last_exc}"
                    break

                _chunk_est_tokens.append(est_tok)
                _chunk_actual_tokens.append(
                    _chunk_actual_usage.get("prompt_tokens") if _chunk_actual_usage else None
                )
                _chunk_completion_tokens.append(
                    _chunk_actual_usage.get("completion_tokens") if _chunk_actual_usage else None
                )
                chunk_results[str(idx)] = chunk_parsed
                completed_chunks.add(idx)
                chunk_state = {
                    "total_chunks": total_chunks,
                    "chunk_word_budget": chunk_word_budget,
                    "completed": list(completed_chunks),
                    "results": chunk_results,
                }
                data["_llm_chunks"] = chunk_state
                record = db.session.get(SubmissionRecord, record_id)
                if record is None:
                    raise Exception(
                        f"submit_to_llm: SubmissionRecord #{record_id} not found on reload (chunk {idx + 1})"
                    )
                record.set_language_analysis_data(data)
                try:
                    db.session.commit()
                except SQLAlchemyError as exc:
                    db.session.rollback()
                    current_app.logger.exception(
                        f"SQLAlchemyError committing chunk {idx + 1} result",
                        exc_info=exc,
                    )
                    raise self.retry()
                db.session.close()

            if chunk_failed:
                # Record failure and return; intermediate state is preserved in
                # data["_llm_chunks"] so a subsequent re-trigger can resume.
                elapsed = round(time.monotonic() - _t_llm, 1)
                data.setdefault("timings", {})["llm_s"] = elapsed
                # Reload record — the session was closed after the last successful chunk commit.
                record = db.session.get(SubmissionRecord, record_id)
                if record is None:
                    raise Exception(
                        f"submit_to_llm: SubmissionRecord #{record_id} not found on reload (chunk failure)"
                    )
                record.llm_analysis_failed = True
                record.llm_failure_reason = chunk_failure_reason
                if accumulated:
                    data["llm_raw_response"] = accumulated
                errors.append(
                    {
                        "stage": "llm_submission",
                        "type": type(last_exc).__name__ if last_exc else "ChunkFailure",
                        "message": chunk_failure_reason,
                    }
                )
                current_app.logger.error(
                    f"language_analysis.submit_to_llm: {chunk_failure_reason} "
                    f"for record #{record_id}"
                )
                data["errors"] = errors
                record.set_language_analysis_data(data)
                try:
                    db.session.commit()
                except SQLAlchemyError as exc:
                    db.session.rollback()
                    current_app.logger.exception(
                        "SQLAlchemyError committing chunk failure", exc_info=exc
                    )
                    raise self.retry()
                return

            # -- Step 3: synthesis (reduce phase) -------------------------
            merged = _merge_chunk_evidence(chunk_results, all_criterion_codes)

            # Override aggregated metadata with the dedicated extraction result,
            # which is more reliable (regex-located, purpose-built prompt).
            merged["metadata"] = {
                "stated_word_count_found": metadata_result.get(
                    "stated_word_count_found", False
                ),
                "stated_word_count": metadata_result.get("stated_word_count"),
                "genai_statement_found": metadata_result.get(
                    "genai_statement_found", False
                ),
                "genai_statement": metadata_result.get("genai_statement", ""),
                "preface_found": metadata_result.get("preface_found", False),
                "preface_precis": metadata_result.get("preface_precis", ""),
            }

            evidence_text = _build_synthesis_evidence_text(
                merged, total_chunks, rubric_snap
            )

            _system_prompt = _build_system_prompt(False, rubric_snap)
            prompt_hash_val = _prompt_hash(_system_prompt)
            parsed_result, accumulated, last_exc, est_tok, _ = _call_llm(
                base_url,
                model,
                _system_prompt,
                _build_synthesis_user_prompt(evidence_text),
                llm_response_schema,
                options={"num_ctx": max(context_size, _SYNTHESIS_MIN_CTX)},
                validate_fn=_validate_llm_response,
                label=f"submit_to_llm/synthesis (record #{record_id})",
            )

            # Inject metadata from the dedicated extraction call into the synthesis
            # result; the grading schema no longer includes these fields.
            if parsed_result is not None:
                for field in (
                    "stated_word_count_found",
                    "stated_word_count",
                    "genai_statement_found",
                    "genai_statement",
                    "preface_found",
                    "preface_precis",
                ):
                    parsed_result[field] = metadata_result.get(field)

            # Token instrumentation for chunked path.
            if _chunk_est_tokens:
                _total_est_tokens = sum(_chunk_est_tokens)
            _actual_available = [t for t in _chunk_actual_tokens if t is not None]
            if _actual_available:
                _peak_prompt_tokens = max(_actual_available)
                _peak_context_pressure = _peak_prompt_tokens / context_size
                _total_actual_prompt_tokens = sum(_actual_available)
            _ct_available = [t for t in _chunk_completion_tokens if t is not None]
            if _ct_available:
                _peak_completion_tokens = max(_ct_available)
                _total_completion_tokens = sum(_ct_available)

        # ----------------------------------------------------------------
        # Common outcome handling (single-pass and chunked-synthesis paths).
        # ----------------------------------------------------------------
        data.setdefault("timings", {})["llm_s"] = round(time.monotonic() - _t_llm, 1)

        # Store LLM provenance in the JSON blob and on the model columns so
        # past runs can be compared if a larger model or wider context window
        # is deployed later.
        data["llm_meta"] = {
            "model": model,
            "context_size": context_size,
            "num_chunks": num_chunks,
        }

        # Merge NLL values (computed before the grading paths) into metrics.
        metrics = data.get("metrics", {})
        metrics["mean_nll"] = mean_nll
        metrics["nll_cv"] = nll_cv  # informational; not a Mahalanobis feature
        data["metrics"] = metrics

        # Reload the record with a fresh DB connection before writing results.
        # The session was closed before the LLM call to prevent connection staleness.
        record = db.session.get(SubmissionRecord, record_id)
        if record is None:
            raise Exception(
                f"submit_to_llm: SubmissionRecord #{record_id} not found on final reload"
            )
        record.llm_model_name = model
        record.llm_context_size = context_size
        record.llm_num_chunks = num_chunks

        # Re-evaluate the AI concern flag now that NLL is available, so full
        # (4D) calibrations can be applied alongside lexical ones.
        try:
            _pclass = (
                record.period.config.project_class
                if record.period and record.period.config
                else None
            )
            _tenant = _pclass.tenant if _pclass else None
            _calibrations = list(_tenant.ai_calibrations) if _tenant else []
        except Exception:
            _calibrations = []

        _ai_result = _ai_concern_flag(
            metrics.get("mattr"),
            metrics.get("mtld"),
            metrics.get("sentence_cv"),
            _calibrations,
            mean_nll=mean_nll,
            nll_cv=metrics.get("nll_cv"),
            llm_model_name=model,
            llm_context_window=context_size,
        )
        _flags = data.get("flags", {})
        _flags["ai_concern"] = _ai_result["concern"]
        _flags["mahalanobis_sigma"] = _ai_result["sigma"]
        _flags["mahalanobis_pvalue"] = _ai_result["p_value"]
        _flags["calibration_results"] = _ai_result.get("calibration_results", [])
        _flags["bonferroni_k"] = _ai_result.get("bonferroni_k", 0)
        _flags["bonferroni_alpha_medium"] = _ai_result.get("bonferroni_alpha_medium")
        _flags["bonferroni_alpha_high"] = _ai_result.get("bonferroni_alpha_high")
        data["flags"] = _flags

        if parsed_result is not None:
            # Success: store result.
            data["llm_result"] = parsed_result
            data["prompt_version"] = PROMPT_VERSION
            data["prompt_hash"] = prompt_hash_val
            # Build a code→criterion mapping so the template can display criterion text.
            # Codes match those used as JSON Schema property keys in _make_band_schema.
            data["criterion_map"] = {
                f"{band_idx}.{crit_idx}": {
                    "text": c["text"],
                    "criterion_id": c["id"],
                    "band_id": band["id"],
                    "rubric_id": rubric_snap.id,
                    "rubric_label": rubric_snap.label,
                }
                for band_idx, band in enumerate(rubric_snap._bands, start=1)
                for crit_idx, c in enumerate(band["criteria"], start=1)
            }
            # Clear any stale feedback from a previous grading run; feedback must
            # re-run against the fresh grading result.
            data.pop("llm_feedback", None)
            # Clean up intermediate chunking state.
            data.pop("_llm_chunks", None)
            data.pop("_llm_metadata", None)
            record.llm_grading_present = True
            record.llm_prompt_version = PROMPT_VERSION
            record.llm_feedback_present = False
            record.llm_feedback_prompt_version = None
        else:
            failure_reason = (
                f"{last_exc} (~{est_tok} est. input tokens)"
                if last_exc
                else "Unknown error"
            )
            record.llm_analysis_failed = True
            record.llm_failure_reason = failure_reason
            record_step_end(_r, record_id, "submit_to_llm", _t0, error=failure_reason)
            data["llm_raw_response"] = accumulated
            errors.append(
                {
                    "stage": "llm_submission",
                    "type": type(last_exc).__name__ if last_exc else "Unknown",
                    "message": failure_reason,
                }
            )
            current_app.logger.error(
                f"language_analysis.submit_to_llm: LLM submission failed for "
                f"record #{record_id}: {failure_reason}"
            )

        data["errors"] = errors
        record.set_language_analysis_data(data)

        try:
            db.session.commit()
        except SQLAlchemyError as exc:
            db.session.rollback()
            current_app.logger.exception(
                "SQLAlchemyError committing LLM result", exc_info=exc
            )
            raise self.retry()

        if not record.llm_analysis_failed:
            record_step_end(
                _r,
                record_id,
                "submit_to_llm",
                _t0,
                meta={
                    "num_chunks": num_chunks,
                    "peak_prompt_tokens": _peak_prompt_tokens,
                    "peak_context_pressure": _peak_context_pressure,
                    "total_est_tokens": _total_est_tokens,
                    "total_actual_prompt_tokens": _total_actual_prompt_tokens,
                    "peak_completion_tokens": _peak_completion_tokens,
                    "total_completion_tokens": _total_completion_tokens,
                },
            )

    # -----------------------------------------------------------------------

    @celery.task(
        bind=True, default_retry_delay=30, soft_time_limit=7200, time_limit=7260
    )
    def submit_to_llm_feedback(self, record_id: int):
        """
        Stage 4 (llm_tasks queue): submit the extracted report text to the LLM
        for formative feedback generation (positive feedback and improvement
        suggestions).  Runs after submit_to_llm in the chain; both tasks read from
        the shared MongoDB scraped-text cache without redundant downloads.

        If the document exceeds the per-chunk feedback word budget, it is split
        into chunks via _build_chunks() and each chunk is submitted independently.
        Results from all successful chunks are merged using _deduplicate_feedback().
        At least one chunk must succeed for feedback to be stored; partial success
        (some chunks succeed, others fail) is accepted — the task is non-fatal.

        Failures are recorded in the JSON blob and in llm_feedback_failed on the
        record, but do not abort the chain — finalize() still runs so the rest of
        the analysis results remain available.
        """
        self.update_state(state=states.STARTED, meta={"msg": "Generating feedback"})

        _r = None
        try:
            _r = get_pipeline_redis()
        except Exception:
            pass
        _t0 = record_step_start(_r, record_id, "submit_to_llm_feedback")

        try:
            record: SubmissionRecord = (
                db.session.query(SubmissionRecord).filter_by(id=record_id).first()
            )
        except SQLAlchemyError as exc:
            current_app.logger.exception(
                "SQLAlchemyError in submit_to_llm_feedback", exc_info=exc
            )
            raise self.retry()

        if record is None:
            raise Exception(
                f"language_analysis.submit_to_llm_feedback: SubmissionRecord #{record_id} not found"
            )

        # Do not re-attempt if a human administrator has not yet cleared the failure flag
        if record.llm_feedback_failed:
            current_app.logger.info(
                f"language_analysis.submit_to_llm_feedback: skipping record #{record_id} — llm_feedback_failed is set"
            )
            record_step_end(_r, record_id, "submit_to_llm_feedback", _t0)
            return

        # Idempotency check: skip if feedback is already present at the current prompt version.
        # Note: submit_to_llm clears llm_feedback_present whenever grading re-runs, ensuring
        # feedback is always regenerated after a grading change.
        if (
            record.llm_feedback_present
            and record.llm_feedback_prompt_version == PROMPT_VERSION
        ):
            current_app.logger.info(
                f"language_analysis.submit_to_llm_feedback: skipping record #{record_id} — "
                f"feedback already present at prompt_version={PROMPT_VERSION}"
            )
            record_step_end(_r, record_id, "submit_to_llm_feedback", _t0)
            return

        data = record.language_analysis_data
        _cached = get_scraped_text(record_id)
        raw_text: str = _cached["scraped_text"] if _cached else ""

        # Build submission text using the 3-tier strategy:
        #   Tier 1: core + appendices if they fit within the context window.
        #   Tier 2: core only (appendix stripped) if that fits.
        #   Tier 3: core chunked if it still exceeds the budget.
        # The reference list is never submitted (same rationale as submit_to_llm).
        _core, _, _appendices = _split_document(raw_text)
        clean_core = _strip_math_lines(_core)
        clean_full = (
            clean_core + "\n\n" + _strip_math_lines(_appendices)
            if _appendices
            else clean_core
        )

        context_size: int = current_app.config.get("OLLAMA_CONTEXT_SIZE", 18432)
        base_url: str = current_app.config.get(
            "OLLAMA_BASE_URL", "http://localhost:11434"
        )
        model: str = current_app.config.get("OLLAMA_MODEL", "llama3.1:70b")

        # Release the DB connection before the long-running LLM call (same rationale
        # as submit_to_llm).  The record will be reloaded before the final write.
        db.session.close()

        _t_feedback = time.monotonic()

        feedback_word_budget = max(
            int(
                (context_size - _FEEDBACK_OVERHEAD_TOKENS)
                / _TOKENS_PER_WORD_CONTENT
                * 0.90
            ),
            500,
        )
        full_words = len(clean_full.split())
        core_words = len(clean_core.split())

        if full_words <= feedback_word_budget:
            # Tier 1: core + appendices fits within the context window.
            tier = 1
            chunk_texts = [clean_full]
        elif core_words <= feedback_word_budget:
            # Tier 2: appendix stripped; core fits without chunking.
            tier = 2
            chunk_texts = [clean_core]
            current_app.logger.info(
                f"language_analysis.submit_to_llm_feedback: record #{record_id} — "
                f"appendix excluded to fit context window ({full_words} → {core_words} words)"
            )
        else:
            # Tier 3: core still exceeds budget; chunk it.
            tier = 3
            chunk_texts = _build_chunks(clean_core, feedback_word_budget)
            if not chunk_texts:
                chunk_texts = [clean_core]
            current_app.logger.info(
                f"language_analysis.submit_to_llm_feedback: record #{record_id} — "
                f"{core_words} core words split into {len(chunk_texts)} chunk(s) "
                f"(full_words={full_words})"
            )

        feedback_results: list[dict] = []
        last_exc: Exception | None = None
        est_tok: int = 0
        _fb_chunk_est_tokens: list[int] = []
        _fb_chunk_actual_tokens: list[int | None] = []
        _fb_chunk_completion_tokens: list[int | None] = []

        for chunk_idx, chunk_text in enumerate(chunk_texts):
            chunk_parsed, _, chunk_exc, est_tok, _fb_actual_usage = _call_llm(
                base_url,
                model,
                _build_feedback_system_prompt(False),
                _build_feedback_user_prompt(chunk_text),
                _LLM_FEEDBACK_RESPONSE_SCHEMA,
                options={"num_ctx": context_size},
                validate_fn=_validate_feedback_response,
                label=(
                    f"submit_to_llm_feedback/chunk {chunk_idx + 1}/{len(chunk_texts)} "
                    f"(record #{record_id})"
                ),
                user_tokens_per_word=_TOKENS_PER_WORD_CONTENT,
            )
            _fb_chunk_est_tokens.append(est_tok)
            _fb_chunk_actual_tokens.append(
                _fb_actual_usage.get("prompt_tokens") if _fb_actual_usage else None
            )
            _fb_chunk_completion_tokens.append(
                _fb_actual_usage.get("completion_tokens") if _fb_actual_usage else None
            )
            if chunk_parsed is not None:
                feedback_results.append(chunk_parsed)
            else:
                last_exc = chunk_exc
                current_app.logger.warning(
                    f"language_analysis.submit_to_llm_feedback: chunk {chunk_idx + 1}/"
                    f"{len(chunk_texts)} failed for record #{record_id}: {chunk_exc}"
                )

        _fb_actual_available = [t for t in _fb_chunk_actual_tokens if t is not None]
        _fb_peak_prompt_tokens: int | None = max(_fb_actual_available) if _fb_actual_available else None
        _fb_peak_context_pressure: float | None = (
            _fb_peak_prompt_tokens / context_size if _fb_peak_prompt_tokens is not None else None
        )
        _fb_total_est_tokens: int | None = sum(_fb_chunk_est_tokens) if _fb_chunk_est_tokens else None
        _fb_total_actual_prompt_tokens: int | None = (
            sum(_fb_actual_available) if _fb_actual_available else None
        )
        _fb_ct_available = [t for t in _fb_chunk_completion_tokens if t is not None]
        _fb_peak_completion_tokens: int | None = max(_fb_ct_available) if _fb_ct_available else None
        _fb_total_completion_tokens: int | None = sum(_fb_ct_available) if _fb_ct_available else None

        data.setdefault("timings", {})["llm_feedback_s"] = round(
            time.monotonic() - _t_feedback, 1
        )
        errors: list = data.get("errors", [])

        # Reload the record with a fresh DB connection before writing results.
        record = db.session.get(SubmissionRecord, record_id)
        if record is None:
            raise Exception(
                f"submit_to_llm_feedback: SubmissionRecord #{record_id} not found on reload"
            )

        if feedback_results:
            # Merge all successful chunk results.
            all_positive: list[str] = []
            all_improvements: list[str] = []
            for r in feedback_results:
                all_positive.extend(r.get("positive_feedback", []))
                all_improvements.extend(r.get("improvements", []))
            data["llm_feedback"] = {
                "positive_feedback": _deduplicate_feedback(all_positive, 3),
                "improvements": _deduplicate_feedback(all_improvements, 3),
            }
            # Explicitly mark success (False = ran and succeeded).  This completes the
            # three-way semantics: None = not yet run, False = succeeded, True = failed.
            record.llm_feedback_failed = False
            record.llm_feedback_failure_reason = None
            record.llm_feedback_present = True
            record.llm_feedback_prompt_version = PROMPT_VERSION
        else:
            failure_reason = (
                f"{last_exc} (~{est_tok} est. input tokens)"
                if last_exc
                else "Unknown error"
            )
            record.llm_feedback_failed = True
            record.llm_feedback_failure_reason = failure_reason
            errors.append(
                {
                    "stage": "llm_feedback_submission",
                    "type": type(last_exc).__name__ if last_exc else "Unknown",
                    "message": failure_reason,
                }
            )
            current_app.logger.error(
                f"language_analysis.submit_to_llm_feedback: all chunks failed for "
                f"record #{record_id}: {failure_reason}"
            )

        data["errors"] = errors
        record.set_language_analysis_data(data)

        try:
            db.session.commit()
        except SQLAlchemyError as exc:
            db.session.rollback()
            current_app.logger.exception(
                "SQLAlchemyError committing feedback result", exc_info=exc
            )
            raise self.retry()

        _fb_meta = {
            "num_chunks": len(chunk_texts),
            "peak_prompt_tokens": _fb_peak_prompt_tokens,
            "peak_context_pressure": _fb_peak_context_pressure,
            "total_est_tokens": _fb_total_est_tokens,
            "total_actual_prompt_tokens": _fb_total_actual_prompt_tokens,
            "tier": tier,
            "feedback_word_budget": feedback_word_budget,
            "peak_completion_tokens": _fb_peak_completion_tokens,
            "total_completion_tokens": _fb_total_completion_tokens,
        }
        if record.llm_feedback_failed:
            record_step_end(
                _r,
                record_id,
                "submit_to_llm_feedback",
                _t0,
                error=record.llm_feedback_failure_reason or "Feedback generation failed",
                meta=_fb_meta,
            )
        else:
            record_step_end(_r, record_id, "submit_to_llm_feedback", _t0, meta=_fb_meta)

    # -----------------------------------------------------------------------

    @celery.task(bind=True, default_retry_delay=30)
    def finalize_language_step(self, record_id: int):
        """
        Stage 5 (default queue): mark language analysis as complete and reset any prior
        chunking-failure state.  Risk factors are NOT computed here — they are computed
        in finalize_risk_flags after the similarity scan has run, so that
        risk_factors.similarity_flagged reflects fresh data.
        """
        _r = None
        try:
            _r = get_pipeline_redis()
        except Exception:
            pass
        _t0 = record_step_start(_r, record_id, "finalize_language_step")

        try:
            record: SubmissionRecord = (
                db.session.query(SubmissionRecord).filter_by(id=record_id).first()
            )
        except SQLAlchemyError as exc:
            record_step_end(
                _r, record_id, "finalize_language_step", _t0, error=repr(exc)
            )
            current_app.logger.exception(
                "SQLAlchemyError in language_analysis.finalize_language_step",
                exc_info=exc,
            )
            raise self.retry()

        if record is None:
            record_step_end(
                _r,
                record_id,
                "finalize_language_step",
                _t0,
                error="SubmissionRecord not found",
            )
            raise Exception(
                f"language_analysis.finalize_language_step: SubmissionRecord #{record_id} not found"
            )

        record.language_analysis_complete = True

        # Clear any chunking-failure state from a previous attempt so a successful
        # retry does not leave stale failure flags visible to reviewers.
        record.llm_chunking_failed = False
        record.llm_chunking_failure_reason = None

        try:
            db.session.commit()
        except SQLAlchemyError as exc:
            db.session.rollback()
            record_step_end(
                _r, record_id, "finalize_language_step", _t0, error=repr(exc)
            )
            current_app.logger.exception(
                "SQLAlchemyError in language_analysis.finalize_language_step commit",
                exc_info=exc,
            )
            raise self.retry()

        record_step_end(_r, record_id, "finalize_language_step", _t0)

    # -----------------------------------------------------------------------

    @celery.task(bind=True, default_retry_delay=30)
    def finalize_risk_flags(self, record_id: int):
        """
        Stage 9 (default queue): compute risk factors and trigger processed-report
        generation.  Runs after run_similarity_check so that risk_factors.similarity_flagged
        reflects the freshest SimilarityConcern data, and the processed-report template
        can incorporate similarity results if desired.
        """
        _r = None
        try:
            _r = get_pipeline_redis()
        except Exception:
            pass
        _t0 = record_step_start(_r, record_id, "finalize_risk_flags")

        try:
            record: SubmissionRecord = (
                db.session.query(SubmissionRecord).filter_by(id=record_id).first()
            )
        except SQLAlchemyError as exc:
            record_step_end(_r, record_id, "finalize_risk_flags", _t0, error=repr(exc))
            current_app.logger.exception(
                "SQLAlchemyError in language_analysis.finalize_risk_flags", exc_info=exc
            )
            raise self.retry()

        if record is None:
            record_step_end(
                _r,
                record_id,
                "finalize_risk_flags",
                _t0,
                error="SubmissionRecord not found",
            )
            raise Exception(
                f"language_analysis.finalize_risk_flags: SubmissionRecord #{record_id} not found"
            )

        # Compute/refresh risk factors using current analysis data and project configuration.
        try:
            config = record.period.config if record.period else None
            record.compute_risk_factors(config)
        except Exception as exc:
            # Non-fatal: log and continue — analysis results are still available.
            current_app.logger.exception(
                "Exception computing risk factors in language_analysis.finalize_risk_flags",
                exc_info=exc,
            )

        try:
            log_db_commit(
                "Language analysis workflow completed",
                student=record.owner.student if record.owner else None,
                project_classes=record.owner.config.project_class
                if record.owner and record.owner.config
                else None,
                endpoint=self.name,
            )
        except SQLAlchemyError as exc:
            db.session.rollback()
            current_app.logger.exception(
                "SQLAlchemyError in language_analysis.finalize_risk_flags commit",
                exc_info=exc,
            )
            raise self.retry()

        # Trigger (re-)generation of the processed report now that both LLM and
        # similarity data are available.
        if record.report is not None:
            _dispatch_process_report(record_id)

        record_step_end(_r, record_id, "finalize_risk_flags", _t0)

    # -----------------------------------------------------------------------

    @celery.task(bind=True, default_retry_delay=30)
    def error_handler(self, record_id: int, user_id: int):
        """
        Error callback: called when any task in the chain raises an unhandled
        exception.  Marks the analysis as not started (so it can be re-triggered)
        and logs the failure.

        Note: llm_analysis_failed is NOT set here — it is reserved for LLM
        inference and JSON parsing failures only.
        """
        try:
            record: SubmissionRecord = (
                db.session.query(SubmissionRecord).filter_by(id=record_id).first()
            )
        except SQLAlchemyError as exc:
            current_app.logger.exception(
                "SQLAlchemyError in language_analysis.error_handler", exc_info=exc
            )
            return

        if record is None:
            current_app.logger.error(
                f"language_analysis.error_handler: SubmissionRecord #{record_id} not found"
            )
            return

        # Reset progress flags so the analysis can be re-triggered
        record.language_analysis_started = False
        record.language_analysis_complete = False

        # Record the workflow-level failure in the JSON blob
        data = record.language_analysis_data
        data.setdefault("errors", []).append(
            {
                "stage": "workflow",
                "type": "UnhandledError",
                "message": "An unhandled exception occurred in the language analysis workflow. Check Celery logs.",
            }
        )
        record.set_language_analysis_data(data)

        try:
            db.session.commit()
        except SQLAlchemyError as exc:
            db.session.rollback()
            current_app.logger.exception(
                "SQLAlchemyError in language_analysis.error_handler commit",
                exc_info=exc,
            )

    # ---------------------------------------------------------------------------
    # Helpers shared by both recalculate modes
    # ---------------------------------------------------------------------------

    def _reclassify_record(record, calibrations: list) -> bool:
        """Re-run _ai_concern_flag and risk factors for *record* using stored metrics.

        Returns True on success.  Caller is responsible for committing.
        """
        la = record.language_analysis_data
        metrics = la.get("metrics", {})
        flags = la.get("flags", {})

        ai_result = _ai_concern_flag(
            metrics.get("mattr"),
            metrics.get("mtld"),
            metrics.get("sentence_cv"),
            calibrations,
            mean_nll=metrics.get("mean_nll"),
            nll_cv=metrics.get("nll_cv"),
            llm_model_name=record.llm_model_name,
            llm_context_window=record.llm_context_size,
        )
        flags["ai_concern"] = ai_result["concern"]
        flags["mahalanobis_sigma"] = ai_result["sigma"]
        flags["mahalanobis_pvalue"] = ai_result["p_value"]
        flags["calibration_results"] = ai_result.get("calibration_results", [])
        flags["bonferroni_k"] = ai_result.get("bonferroni_k", 0)
        flags["bonferroni_alpha_medium"] = ai_result.get("bonferroni_alpha_medium")
        flags["bonferroni_alpha_high"] = ai_result.get("bonferroni_alpha_high")
        la["flags"] = flags
        record.set_language_analysis_data(la)

        try:
            config = record.period.config if record.period else None
            record.compute_risk_factors(config)
        except Exception as exc:
            current_app.logger.warning(
                f"recalculate_ai_concern: could not recompute risk factors for record #{record.id}: {exc}"
            )
        return True

    # ---------------------------------------------------------------------------
    # Fan-out sub-task: process one pclass×year batch for full metric recompute
    # ---------------------------------------------------------------------------

    @celery.task(bind=True, default_retry_delay=30)
    def recalculate_ai_concern_batch(
        self, task_id: str, tenant_id: int, record_ids: list
    ):
        """
        Process a batch of SubmissionRecord IDs for full lexical-metrics recomputation.

        For each record the task:
          1. Reads scraped text from the MongoDB scraped-text cache.
          2. Re-runs _split_document, _strip_math_lines.
          3. Recomputes MATTR, MTLD, burstiness, and sentence CV via the current
             pipeline implementations (including any code-block filtering).
          4. Updates metric values and classification flags in the JSON blob.
          5. Re-evaluates the Mahalanobis AI concern flag.
          6. Re-runs compute_risk_factors().
          7. Commits in batches of 50.

        On a MongoDB cache miss the report asset is re-downloaded from the object
        store and the result is stored in the cache for future use.
        Records with no report asset or a missing object-store file are skipped.
        Per-record errors are caught so that the chord callback always fires.

        Returns {"updated": int, "skipped": int, "errors": int}.
        """
        updated = skipped = errors = 0

        try:
            tenant: Tenant = db.session.query(Tenant).filter_by(id=tenant_id).first()
        except SQLAlchemyError as exc:
            current_app.logger.exception(
                f"recalculate_ai_concern_batch: DB error loading tenant #{tenant_id}",
                exc_info=exc,
            )
            return {"updated": 0, "skipped": len(record_ids), "errors": 0}

        if tenant is None:
            current_app.logger.error(
                f"recalculate_ai_concern_batch: Tenant #{tenant_id} not found"
            )
            return {"updated": 0, "skipped": len(record_ids), "errors": 0}

        calibrations = list(tenant.ai_calibrations)
        if not calibrations:
            current_app.logger.warning(
                "recalculate_ai_concern_batch: no calibration data — skipping batch"
            )
            return {"updated": 0, "skipped": len(record_ids), "errors": 0}

        for i, record_id in enumerate(record_ids, start=1):
            try:
                record = (
                    db.session.query(SubmissionRecord).filter_by(id=record_id).first()
                )
                if record is None:
                    current_app.logger.warning(
                        f"recalculate_ai_concern_batch: record #{record_id} not found"
                    )
                    skipped += 1
                    continue

                la = record.language_analysis_data
                _cached = get_scraped_text(record_id)
                raw_text = _cached["scraped_text"] if _cached else None

                if not raw_text:
                    # Cache miss — re-download from object store and populate the MongoDB cache.
                    if record.report is None:
                        current_app.logger.warning(
                            f"recalculate_ai_concern_batch: record #{record_id} has no report asset — skipping"
                        )
                        skipped += 1
                        continue

                    asset = record.report
                    storage = current_app.config["OBJECT_STORAGE_ASSETS"]
                    adapter = AssetCloudAdapter(
                        asset,
                        storage,
                        audit_data=f"recalculate_ai_concern_batch (record #{record_id})",
                    )
                    if not adapter.exists():
                        current_app.logger.warning(
                            f"recalculate_ai_concern_batch: report asset not found in object store for record #{record_id} — skipping"
                        )
                        skipped += 1
                        continue

                    mimetype = (asset.mimetype or "").lower()
                    page_count = 0
                    with adapter.download_to_scratch() as scratch:
                        path = str(scratch.path)
                        if "pdf" in mimetype or path.lower().endswith(".pdf"):
                            raw_text, page_count = _extract_pdf_text(path)
                        elif (
                            "word" in mimetype
                            or "officedocument" in mimetype
                            or path.lower().endswith((".docx", ".doc"))
                        ):
                            raw_text, page_count = _extract_docx_text(path)
                        else:
                            raw_text, page_count = _extract_pdf_text(path)
                    raw_text = unicodedata.normalize("NFKC", raw_text)
                    store_scraped_text(
                        record_id, asset.id, mimetype, raw_text, page_count
                    )

                # Re-process from text using the current pipeline.
                _core, _references, _appendices = _split_document(raw_text)
                content_text = (_core + "\n\n" + _appendices) if _appendices else _core
                clean_content = _strip_math_lines(content_text)

                mattr, mtld = _compute_mattr_mtld(clean_content)
                burstiness_groups, burstiness = _compute_burstiness(raw_text)
                sentence_cv = _compute_sentence_cv(clean_content)

                metrics = la.get("metrics", {})
                metrics["mattr"] = mattr
                metrics["mtld"] = mtld
                metrics["burstiness"] = burstiness
                metrics["burstiness_by_group"] = burstiness_groups
                metrics["sentence_cv"] = sentence_cv
                metrics["mattr_flag"] = classify_mattr(mattr)
                metrics["mtld_flag"] = classify_mtld(mtld)
                metrics["burstiness_flag"] = classify_burstiness(burstiness)
                metrics["sentence_cv_flag"] = classify_sentence_cv(sentence_cv)
                la["metrics"] = metrics

                # Re-classify using fresh metric values.
                flags = la.get("flags", {})
                ai_result = _ai_concern_flag(
                    mattr,
                    mtld,
                    sentence_cv,
                    calibrations,
                    mean_nll=metrics.get("mean_nll"),
                    llm_model_name=record.llm_model_name,
                    llm_context_window=record.llm_context_size,
                )
                flags["ai_concern"] = ai_result["concern"]
                flags["mahalanobis_sigma"] = ai_result["sigma"]
                flags["mahalanobis_pvalue"] = ai_result["p_value"]
                flags["calibration_results"] = ai_result.get("calibration_results", [])
                flags["bonferroni_k"] = ai_result.get("bonferroni_k", 0)
                flags["bonferroni_alpha_medium"] = ai_result.get(
                    "bonferroni_alpha_medium"
                )
                flags["bonferroni_alpha_high"] = ai_result.get("bonferroni_alpha_high")
                la["flags"] = flags
                record.set_language_analysis_data(la)

                try:
                    config = record.period.config if record.period else None
                    record.compute_risk_factors(config)
                except Exception as exc:
                    current_app.logger.warning(
                        f"recalculate_ai_concern_batch: could not recompute risk factors for record #{record.id}: {exc}"
                    )

                updated += 1

                if updated % 50 == 0:
                    try:
                        db.session.commit()
                    except SQLAlchemyError as exc:
                        db.session.rollback()
                        current_app.logger.exception(
                            "recalculate_ai_concern_batch: DB error during batch commit",
                            exc_info=exc,
                        )

            except Exception as exc:
                current_app.logger.warning(
                    f"recalculate_ai_concern_batch: error on record #{record_id}: {exc}"
                )
                errors += 1

        try:
            db.session.commit()
        except SQLAlchemyError as exc:
            db.session.rollback()
            current_app.logger.exception(
                "recalculate_ai_concern_batch: final commit error", exc_info=exc
            )

        return {"updated": updated, "skipped": skipped, "errors": errors}

    # ---------------------------------------------------------------------------
    # Chord finalize / error callbacks for fan-out full recalculation
    # ---------------------------------------------------------------------------

    @celery.task(bind=False)
    def recalculate_ai_concern_finalize(results, task_id: str, total: int):
        """
        Chord callback: aggregate per-batch results and mark the TaskRecord SUCCESS.

        *results* is the list of {"updated": int, "skipped": int, "errors": int}
        dicts returned by each recalculate_ai_concern_batch sub-task.
        """
        updated = sum(r.get("updated", 0) for r in results if isinstance(r, dict))
        skipped = sum(r.get("skipped", 0) for r in results if isinstance(r, dict))
        errors = sum(r.get("errors", 0) for r in results if isinstance(r, dict))

        parts = [f"updated {updated}/{total}"]
        if skipped:
            parts.append(f"skipped {skipped} (no cached text)")
        if errors:
            parts.append(f"errors {errors}")
        msg = "Full lexical-metrics recalculation complete — " + ", ".join(parts) + "."

        progress_update(task_id, TaskRecord.SUCCESS, 100, msg, autocommit=True)

    @celery.task(bind=False)
    def recalculate_ai_concern_error(task_id: str):
        """Error callback for the full recalculation chord."""
        progress_update(
            task_id,
            TaskRecord.FAILURE,
            100,
            "Full recalculation failed — see worker logs for details.",
            autocommit=True,
        )

    # ---------------------------------------------------------------------------

    @celery.task(bind=True, default_retry_delay=30)
    def recalculate_ai_concern(
        self,
        task_id: str,
        tenant_id: int,
        pclass_ids=None,
        years=None,
        full_recalculate=False,
    ):
        """
        Re-evaluate the Mahalanobis-based AI concern flag for all completed
        SubmissionRecords belonging to *tenant_id*, optionally filtered to
        specific project class IDs and/or academic years.

        When *full_recalculate* is False (default) the task re-runs only the
        Mahalanobis classification from already-stored metric values — fast,
        DB-only, sequential.

        When *full_recalculate* is True the task re-processes the cached
        extracted text through the current metric pipeline (MATTR, MTLD,
        burstiness, sentence CV) before reclassifying.  Work is distributed
        across Celery workers by fanning out one sub-task per (pclass × year)
        batch via a Celery chord.  Records without cached extracted text are
        skipped.
        """
        self.update_state(
            state="STARTED", meta={"msg": "Preparing AI concern recalculation"}
        )
        progress_update(task_id, TaskRecord.RUNNING, 5, "Querying submissions…")

        # Fetch tenant and its calibration.
        try:
            tenant: Tenant = db.session.query(Tenant).filter_by(id=tenant_id).first()
        except SQLAlchemyError as exc:
            current_app.logger.exception(
                "recalculate_ai_concern: DB error loading tenant", exc_info=exc
            )
            progress_update(
                task_id,
                TaskRecord.FAILURE,
                100,
                "Database error loading tenant",
                autocommit=True,
            )
            return

        if tenant is None:
            current_app.logger.error(
                f"recalculate_ai_concern: Tenant #{tenant_id} not found"
            )
            progress_update(
                task_id, TaskRecord.FAILURE, 100, "Tenant not found", autocommit=True
            )
            return

        calibrations = list(tenant.ai_calibrations)
        if not calibrations:
            current_app.logger.warning(
                "recalculate_ai_concern: tenant has no calibration data — aborting"
            )
            progress_update(
                task_id,
                TaskRecord.FAILURE,
                100,
                "No calibration data available",
                autocommit=True,
            )
            return

        # Build query for target records, fetching pclass_id and year for grouping.
        try:
            q = (
                db.session.query(
                    SubmissionRecord, ProjectClass.id, ProjectClassConfig.year
                )
                .join(
                    SubmissionPeriodRecord,
                    SubmissionRecord.period_id == SubmissionPeriodRecord.id,
                )
                .join(
                    ProjectClassConfig,
                    SubmissionPeriodRecord.config_id == ProjectClassConfig.id,
                )
                .join(ProjectClass, ProjectClassConfig.pclass_id == ProjectClass.id)
                .filter(ProjectClass.tenant_id == tenant_id)
                .filter(SubmissionRecord.language_analysis_complete == True)  # noqa: E712
            )
            if pclass_ids:
                q = q.filter(ProjectClass.id.in_(pclass_ids))
            if years:
                q = q.filter(ProjectClassConfig.year.in_(years))

            rows = q.all()
        except SQLAlchemyError as exc:
            current_app.logger.exception(
                "recalculate_ai_concern: DB error querying records", exc_info=exc
            )
            progress_update(
                task_id,
                TaskRecord.FAILURE,
                100,
                "Database error querying records",
                autocommit=True,
            )
            return

        total = len(rows)
        if total == 0:
            progress_update(
                task_id,
                TaskRecord.SUCCESS,
                100,
                "No eligible submissions found.",
                autocommit=True,
            )
            return

        # ── Full recalculation: fan out by (pclass × year) ───────────────────
        if full_recalculate:
            # Group record IDs by (pclass_id, year).
            from collections import defaultdict

            batch_map: dict[tuple, list[int]] = defaultdict(list)
            for record, pclass_id, year in rows:
                batch_map[(pclass_id, year)].append(record.id)

            batches = list(batch_map.values())
            n_batches = len(batches)

            progress_update(
                task_id,
                TaskRecord.RUNNING,
                10,
                f"Dispatching full metric recomputation for {total} submission(s) across {n_batches} batch(es)…",
            )

            if n_batches == 1:
                # Only one batch — run inline to avoid chord overhead.
                result = recalculate_ai_concern_batch.run(
                    task_id, tenant_id, batches[0]
                )
                updated = result.get("updated", 0)
                skipped = result.get("skipped", 0)
                errors = result.get("errors", 0)
                parts = [f"updated {updated}/{total}"]
                if skipped:
                    parts.append(f"skipped {skipped} (no cached text)")
                if errors:
                    parts.append(f"errors {errors}")
                progress_update(
                    task_id,
                    TaskRecord.SUCCESS,
                    100,
                    "Full lexical-metrics recalculation complete — "
                    + ", ".join(parts)
                    + ".",
                    autocommit=True,
                )
                return

            # Two or more batches — fan out via chord.
            sub_tasks = cgroup(
                [
                    recalculate_ai_concern_batch.si(task_id, tenant_id, batch_ids)
                    for batch_ids in batches
                ]
            )
            finalize = recalculate_ai_concern_finalize.s(task_id, total)
            error_cb = recalculate_ai_concern_error.si(task_id)
            self.replace(chord(sub_tasks, finalize).on_error(error_cb))
            return

        # ── Classify-only: sequential loop (existing behaviour) ───────────────
        records = [row[0] for row in rows]
        progress_update(
            task_id,
            TaskRecord.RUNNING,
            10,
            f"Recalculating AI concern for {total} submission(s)…",
        )

        updated = 0
        for i, record in enumerate(records, start=1):
            try:
                _reclassify_record(record, calibrations)
                updated += 1

                if updated % 50 == 0:
                    try:
                        db.session.commit()
                    except SQLAlchemyError as exc:
                        db.session.rollback()
                        current_app.logger.exception(
                            "recalculate_ai_concern: DB error during batch commit",
                            exc_info=exc,
                        )

                pct = 10 + int(85 * i / total)
                progress_update(
                    task_id, TaskRecord.RUNNING, pct, f"Processed {i}/{total}…"
                )

            except Exception as exc:
                current_app.logger.warning(
                    f"recalculate_ai_concern: error processing record #{record.id}: {exc}"
                )

        try:
            db.session.commit()
        except SQLAlchemyError as exc:
            db.session.rollback()
            current_app.logger.exception(
                "recalculate_ai_concern: final commit error", exc_info=exc
            )
            progress_update(
                task_id,
                TaskRecord.FAILURE,
                100,
                "Database error on final commit",
                autocommit=True,
            )
            return

        progress_update(
            task_id,
            TaskRecord.SUCCESS,
            100,
            f"AI concern recalculated for {updated} submission(s).",
            autocommit=True,
        )

    return recalculate_ai_concern
